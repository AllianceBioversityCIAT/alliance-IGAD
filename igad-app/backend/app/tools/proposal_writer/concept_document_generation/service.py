"""
Concept Document Generation Service

Generates comprehensive concept documents based on:
- RFP analysis and requirements
- User concept evaluation and selections
- Proposal outline structure and guidance

The service enriches selected sections with proposal outline data
and uses Claude to generate detailed, donor-aligned documentation.
"""

import json
import logging
import os
import re
import traceback
from datetime import datetime
from io import BytesIO
from typing import Any, Dict, Optional

import boto3
from boto3.dynamodb.conditions import Attr
from docx import Document
from PyPDF2 import PdfReader

from app.shared.ai.bedrock_service import BedrockService
from app.tools.proposal_writer.concept_document_generation.config import (
    CONCEPT_DOCUMENT_GENERATION_SETTINGS,
)

logger = logging.getLogger(__name__)
logger.setLevel(logging.INFO)


class ConceptDocumentGenerator:
    """
    Generates concept documents using AI.

    Workflow:
    1. Load prompt template from DynamoDB
    2. Retrieve proposal outline (if not provided)
    3. Prepare context with RFP analysis and concept evaluation
    4. Inject context into prompt
    5. Call Claude via Bedrock
    6. Parse and structure the generated document
    """

    def __init__(self):
        """
        Initialize Bedrock, DynamoDB, and S3 clients.

        Sets up clients for:
        - Bedrock: Claude AI invocation
        - DynamoDB: Proposal data and outline retrieval
        - S3: Initial concept document retrieval
        """
        self.bedrock = BedrockService()
        self.dynamodb = boto3.resource("dynamodb")
        self.s3 = boto3.client("s3")
        self.table_name = CONCEPT_DOCUMENT_GENERATION_SETTINGS.get(
            "table_name", "igad-testing-main-table"
        )
        self.bucket = os.environ.get("PROPOSALS_BUCKET")
        if not self.bucket:
            raise Exception("PROPOSALS_BUCKET environment variable not set")

    def generate_document(
        self,
        proposal_code: str,
        rfp_analysis: Dict[str, Any],
        concept_evaluation: Dict[str, Any],
        proposal_outline: Optional[Dict[str, Any]] = None,
        reference_proposals_analysis: Optional[Dict[str, Any]] = None,
        existing_work_analysis: Optional[Dict[str, Any]] = None,
    ) -> Dict[str, Any]:
        """
        Generate concept document using AI.

        Args:
            proposal_code: Proposal identifier for logging
            rfp_analysis: RFP analysis from step 1
            concept_evaluation: User selections from step 2
            proposal_outline: Optional proposal outline (loads from DB if not provided)
            reference_proposals_analysis: Optional reference proposals analysis
            existing_work_analysis: Optional existing work analysis

        Returns:
            Dict with 'generated_concept_document' and 'sections'

        Raises:
            ValueError: If prompt template not found
            Exception: If document generation fails
        """
        try:
            logger.info(f"📋 Generating document for proposal: {proposal_code}")

            # Step 1: Load prompt template
            logger.info("📝 Loading prompt template...")
            prompt_parts = self._get_prompt_template()

            if not prompt_parts:
                raise ValueError("Prompt template not found in DynamoDB")

            # Step 2: Load proposal outline if needed
            if not proposal_outline:
                logger.info("📥 Loading proposal outline...")
                proposal_outline = self._load_proposal_outline(proposal_code)
                if proposal_outline:
                    outline_count = len(proposal_outline.get("proposal_outline", []))
                    logger.info(f"✅ Loaded {outline_count} outline sections")

            # Step 3: Load initial concept from S3
            logger.info("📥 Loading initial concept from S3...")
            initial_concept = self._get_initial_concept_from_s3(proposal_code)
            if initial_concept:
                logger.info(
                    f"✅ Initial concept loaded: {len(initial_concept)} characters"
                )
            else:
                logger.warning("⚠️  Initial concept not found in S3")

            # Step 4: Prepare context
            logger.info("🔄 Preparing context...")
            context = self._prepare_context(
                rfp_analysis,
                concept_evaluation,
                proposal_outline,
                initial_concept,
                reference_proposals_analysis,
                existing_work_analysis,
            )

            # Step 5: Build final prompt
            user_prompt = self._inject_context(prompt_parts["user_prompt"], context)
            final_prompt = f"{user_prompt}\n\n{prompt_parts['output_format']}".strip()

            # Step 6: Call Bedrock
            logger.info("📡 Calling Bedrock (this may take 3-5 minutes)...")
            start_time = datetime.utcnow()

            ai_response = self.bedrock.invoke_claude(
                system_prompt=prompt_parts["system_prompt"],
                user_prompt=final_prompt,
                max_tokens=CONCEPT_DOCUMENT_GENERATION_SETTINGS.get(
                    "max_tokens", 12000
                ),
                temperature=CONCEPT_DOCUMENT_GENERATION_SETTINGS.get(
                    "temperature", 0.2
                ),
                model_id=CONCEPT_DOCUMENT_GENERATION_SETTINGS["model"],
            )

            elapsed = (datetime.utcnow() - start_time).total_seconds()
            logger.info(f"✅ Response received in {elapsed:.1f} seconds")

            # Step 7: Parse response
            logger.info("📊 Parsing response...")
            document = self._parse_response(ai_response)

            logger.info(f"✅ Document generated successfully")
            return document

        except Exception as e:
            logger.error(f"❌ Document generation failed: {str(e)}")
            traceback.print_exc()
            raise

    # ==================== S3 OPERATIONS ====================

    def _get_initial_concept_from_s3(self, proposal_code: str) -> Optional[str]:
        """
        Retrieve initial concept document from S3.

        Tries in order:
        1. Plain text file (concept_text.txt)
        2. PDF/DOC/DOCX files in initial_concept folder

        Args:
            proposal_code: Proposal code for path construction

        Returns:
            Extracted concept text or None if not found

        Raises:
            Returns None if concept not found or extraction fails
        """
        try:
            concept_folder = f"{proposal_code}/documents/initial_concept/"

            # Try text file first (fastest)
            try:
                text_key = f"{concept_folder}concept_text.txt"
                logger.info(f"📄 Trying concept text file: {text_key}")
                obj = self.s3.get_object(Bucket=self.bucket, Key=text_key)
                concept_text = obj["Body"].read().decode("utf-8")
                logger.info(f"✅ Loaded concept text: {len(concept_text)} characters")
                return concept_text
            except Exception as e:
                logger.info(f"ℹ️  No concept_text.txt found: {str(e)}")

            # Try document files (PDF/DOCX)
            logger.info("🔍 Searching for concept document files...")
            response = self.s3.list_objects_v2(
                Bucket=self.bucket, Prefix=concept_folder
            )

            if "Contents" not in response or len(response["Contents"]) == 0:
                logger.warning("⚠️  No concept content found in S3")
                return None

            # Find and extract from first document file
            concept_file = self._find_document_file(response["Contents"])
            if not concept_file:
                logger.warning("⚠️  No document files found (PDF/DOC/DOCX)")
                return None

            logger.info(f"📥 Extracting text from: {concept_file}")
            return self.extract_text_from_file(concept_file)

        except Exception as e:
            logger.error(f"❌ Error loading initial concept: {str(e)}")
            return None

    def _find_document_file(self, contents: list) -> Optional[str]:
        """
        Find first document file in S3 contents list.

        Searches for: PDF, DOC, DOCX files.

        Args:
            contents: List of S3 object summaries

        Returns:
            Document file key if found, None otherwise
        """
        for obj in contents:
            key = obj["Key"]
            if key.lower().endswith((".pdf", ".doc", ".docx")) and not key.endswith(
                "/"
            ):
                return key
        return None

    def extract_text_from_file(self, s3_key: str) -> Optional[str]:
        """
        Extract text from document file (PDF, DOC, or DOCX).

        Args:
            s3_key: S3 key of the file

        Returns:
            Extracted text or None if extraction fails
        """
        try:
            obj = self.s3.get_object(Bucket=self.bucket, Key=s3_key)
            file_bytes = obj["Body"].read()

            if s3_key.lower().endswith(".pdf"):
                return self.extract_text_from_pdf(file_bytes)
            elif s3_key.lower().endswith(".docx"):
                return self.extract_text_from_docx(file_bytes)
            elif s3_key.lower().endswith(".doc"):
                logger.error(
                    "❌ Legacy .doc format not supported. Please upload as .docx or .pdf"
                )
                return None
            else:
                logger.error(f"❌ Unsupported file type: {s3_key}")
                return None

        except Exception as e:
            logger.error(f"❌ File extraction failed: {str(e)}")
            return None

    def extract_text_from_pdf(self, pdf_bytes: bytes) -> str:
        """
        Extract text from PDF bytes.

        Args:
            pdf_bytes: PDF file content as bytes

        Returns:
            Extracted text
        """
        try:
            pdf_file = BytesIO(pdf_bytes)
            reader = PdfReader(pdf_file)
            text = ""

            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"

            return text.strip()

        except Exception as e:
            logger.error(f"❌ PDF extraction failed: {str(e)}")
            return ""

    def extract_text_from_docx(self, docx_bytes: bytes) -> str:
        """
        Extract text from DOCX bytes.

        Extracts from:
        - Paragraphs
        - Tables (with pipe-separated columns)

        Args:
            docx_bytes: DOCX file content as bytes

        Returns:
            Extracted text
        """
        try:
            docx_file = BytesIO(docx_bytes)
            document = Document(docx_file)
            text = ""

            # Extract from paragraphs
            for paragraph in document.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + "\n"

            # Extract from tables
            for table in document.tables:
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells]
                    text += " | ".join(row_text) + "\n"

            return text.strip() if text.strip() else ""

        except Exception as e:
            logger.error(f"❌ DOCX extraction failed: {str(e)}")
            return ""

    # ==================== PRIVATE HELPER METHODS ====================

    def _get_prompt_template(self) -> Optional[Dict[str, str]]:
        """
        Load document generation prompt from DynamoDB.

        Filters for:
        - is_active: true
        - section: "proposal_writer"
        - sub_section: "step-2"
        - categories: contains "Concept Review"

        Returns:
            Dict with 'system_prompt', 'user_prompt', 'output_format', or None
        """
        try:
            table = self.dynamodb.Table(self.table_name)
            response = table.scan(
                FilterExpression=(
                    Attr("is_active").eq(True)
                    & Attr("section").eq("proposal_writer")
                    & Attr("sub_section").eq("step-2")
                    & Attr("categories").contains("Concept Review")
                )
            )

            items = response.get("Items", [])
            if not items:
                logger.warning("⚠️  No prompts found in DynamoDB")
                return None

            prompt_item = items[0]
            logger.info(f"✅ Loaded prompt: {prompt_item.get('name', 'Unnamed')}")

            return {
                "system_prompt": prompt_item.get("system_prompt", ""),
                "user_prompt": prompt_item.get("user_prompt_template", ""),
                "output_format": prompt_item.get("output_format", ""),
            }

        except Exception as e:
            logger.error(f"❌ Error loading prompt: {str(e)}")
            return None

    def _load_proposal_outline(self, proposal_code: str) -> Optional[Dict[str, Any]]:
        """
        Load proposal outline from DynamoDB.

        Args:
            proposal_code: Proposal code to search for

        Returns:
            Proposal outline dict or None
        """
        try:
            table = self.dynamodb.Table(self.table_name)
            response = table.scan(
                FilterExpression=Attr("proposalCode").eq(proposal_code)
            )

            items = response.get("Items", [])
            if not items:
                logger.warning(f"⚠️  No proposal found: {proposal_code}")
                return None

            proposal_outline = items[0].get("proposal_outline")
            if not proposal_outline:
                logger.warning(f"⚠️  No outline found in proposal")
                return None

            logger.info(f"✅ Loaded outline for {proposal_code}")
            return proposal_outline

        except Exception as e:
            logger.error(f"❌ Error loading outline: {str(e)}")
            return None

    def _prepare_context(
        self,
        rfp_analysis: Dict[str, Any],
        concept_evaluation: Dict[str, Any],
        proposal_outline: Optional[Dict[str, Any]] = None,
        initial_concept: Optional[str] = None,
        reference_proposals_analysis: Optional[Dict[str, Any]] = None,
        existing_work_analysis: Optional[Dict[str, Any]] = None,
    ) -> Dict[str, Any]:
        """
        Prepare context for prompt injection.

        Process:
        1. Filter to selected sections
        2. Enrich with outline data
        3. Prepare initial concept (truncate if needed)
        4. Return as JSON/text strings for prompt

        Args:
            rfp_analysis: RFP analysis data
            concept_evaluation: Concept evaluation with selections
            proposal_outline: Optional proposal outline for enrichment
            initial_concept: Optional initial concept document text
            reference_proposals_analysis: Optional reference proposals analysis
            existing_work_analysis: Optional existing work analysis

        Returns:
            Dict with 'rfp_analysis', 'concept_evaluation', 'initial_concept',
            'reference_proposals_analysis', and 'existing_work_analysis'
        """
        # Filter to selected sections
        filtered_evaluation = self._filter_selected_sections(concept_evaluation)

        # Enrich with outline data if available
        if proposal_outline:
            logger.info("🔄 Enriching sections with outline data...")
            enriched_evaluation = self._enrich_with_outline(
                filtered_evaluation, proposal_outline
            )
        else:
            logger.info("ℹ️  Proceeding without outline enrichment")
            enriched_evaluation = filtered_evaluation

        # Prepare initial concept
        prepared_concept = ""
        if initial_concept:
            prepared_concept = self._prepare_initial_concept(initial_concept)
            logger.info(f"📋 Initial concept prepared: {len(prepared_concept)} chars")

        # Build context dict
        context = {
            "rfp_analysis": json.dumps(rfp_analysis, indent=2),
            "concept_evaluation": json.dumps(enriched_evaluation, indent=2),
            "initial_concept": prepared_concept,
        }

        # Add reference proposals analysis if provided
        if reference_proposals_analysis:
            # Unwrap nested structure if present
            analysis = reference_proposals_analysis.get(
                "reference_proposals_analysis", reference_proposals_analysis
            )
            context["reference_proposals_analysis"] = json.dumps(analysis, indent=2)
            logger.info("✅ Reference proposals analysis added to context")

        # Add existing work analysis if provided
        if existing_work_analysis:
            # Unwrap nested structure if present
            analysis = existing_work_analysis.get(
                "existing_work_analysis", existing_work_analysis
            )
            context["existing_work_analysis"] = json.dumps(analysis, indent=2)
            logger.info("✅ Existing work analysis added to context")

        return context

    def _filter_selected_sections(
        self, concept_evaluation: Dict[str, Any]
    ) -> Dict[str, Any]:
        """
        Filter concept evaluation to only selected sections.

        Args:
            concept_evaluation: Full evaluation data

        Returns:
            Filtered evaluation with only selected sections

        Raises:
            Returns original data if filtering fails
        """
        try:
            # Handle nested structure
            concept_analysis = concept_evaluation.get("concept_analysis", {})
            if "concept_analysis" in concept_analysis:
                concept_analysis = concept_analysis["concept_analysis"]

            sections = concept_analysis.get("sections_needing_elaboration", [])
            logger.info(f"📊 Total sections: {len(sections)}")

            # Filter to selected
            selected = [s for s in sections if s.get("selected", False)]
            logger.info(f"✅ Selected: {len(selected)} sections")

            # Log selected titles
            for section in selected:
                title = section.get("section", section.get("title", "Unknown"))
                logger.info(f"   ✓ {title}")

            return {
                "concept_analysis": {
                    **concept_analysis,
                    "sections_needing_elaboration": selected,
                },
                "status": concept_evaluation.get("status", "completed"),
            }

        except Exception as e:
            logger.error(f"❌ Error filtering sections: {str(e)}")
            return concept_evaluation

    def _enrich_with_outline(
        self,
        filtered_evaluation: Dict[str, Any],
        proposal_outline: Dict[str, Any],
    ) -> Dict[str, Any]:
        """
        Enrich selected sections with proposal outline data.

        Adds to each section:
        - recommended_word_count
        - purpose
        - content_guidance
        - guiding_questions

        Args:
            filtered_evaluation: Filtered concept evaluation
            proposal_outline: Proposal outline structure

        Returns:
            Enriched evaluation

        Raises:
            Returns original if enrichment fails
        """
        try:
            outline_sections = proposal_outline.get("proposal_outline", [])
            if not outline_sections:
                logger.warning("⚠️  No sections in proposal_outline")
                return filtered_evaluation

            concept_analysis = filtered_evaluation.get("concept_analysis", {})
            selected_sections = concept_analysis.get("sections_needing_elaboration", [])

            if not selected_sections:
                logger.warning("⚠️  No sections to enrich")
                return filtered_evaluation

            # Create lookup for selected sections only
            selected_titles = [
                s.get("section", s.get("title", "")) for s in selected_sections
            ]
            outline_lookup = {
                os.get("section_title"): os
                for os in outline_sections
                if os.get("section_title") in selected_titles
            }

            logger.info(
                f"📊 Enriching {len(outline_lookup)} sections "
                f"(from {len(outline_sections)} total)"
            )

            # Enrich each section
            enriched_sections = []
            for section in selected_sections:
                section_title = section.get("section", section.get("title", ""))
                outline_data = outline_lookup.get(section_title)

                if outline_data:
                    logger.info(f"   ✓ Enriching {section_title}")
                    content_guidance = outline_data.get("content_guidance", "")

                    # Summarize long guidance
                    if len(content_guidance) > 5000:
                        content_guidance = self._summarize_guidance(content_guidance)

                    enriched_section = {
                        **section,
                        "recommended_word_count": outline_data.get(
                            "recommended_word_count", ""
                        ),
                        "purpose": outline_data.get("purpose", ""),
                        "content_guidance": content_guidance,
                        "guiding_questions": outline_data.get("guiding_questions", []),
                        # Preserve user_comment if it exists
                        "user_comment": section.get("user_comment", ""),
                    }

                    # Log if user comment is present
                    if enriched_section["user_comment"]:
                        logger.info(
                            f"      📝 User comment included ({len(enriched_section['user_comment'])} chars)"
                        )

                    enriched_sections.append(enriched_section)
                else:
                    logger.info(f"   ⚠️  No outline for {section_title}")
                    enriched_sections.append(section)

            logger.info(f"✅ Enriched {len(enriched_sections)} sections")

            return {
                "concept_analysis": {
                    **concept_analysis,
                    "sections_needing_elaboration": enriched_sections,
                },
                "status": filtered_evaluation.get("status", "completed"),
            }

        except Exception as e:
            logger.error(f"❌ Error enriching sections: {str(e)}")
            return filtered_evaluation

    def _prepare_initial_concept(self, concept_text: str) -> str:
        """
        Prepare initial concept for prompt injection.

        Applies intelligent truncation if text exceeds limit:
        - Keeps 60% from beginning (context/problem)
        - Keeps 40% from end (solution/conclusion)
        - Adds truncation notice

        Args:
            concept_text: Full initial concept document text

        Returns:
            Prepared text (possibly truncated)
        """
        max_chars = 50000  # Generous limit for initial concept

        if len(concept_text) <= max_chars:
            logger.info(f"📄 Initial concept within limit ({len(concept_text)} chars)")
            return concept_text

        logger.info(
            f"✂️  Truncating initial concept from {len(concept_text)} to {max_chars} chars"
        )

        chars_to_keep = max_chars - 200
        beginning_chars = int(chars_to_keep * 0.6)
        ending_chars = int(chars_to_keep * 0.4)

        beginning = concept_text[:beginning_chars]
        ending = concept_text[-ending_chars:]

        return (
            f"{beginning}\n\n"
            f"[... Middle section truncated - "
            f"total document: {len(concept_text)} characters ...]\n\n"
            f"{ending}"
        )

    def _summarize_guidance(self, guidance: str) -> str:
        """
        Summarize long content guidance.

        Extracts bullet points or truncates to 500 chars.

        Args:
            guidance: Full content guidance text

        Returns:
            Summarized guidance
        """
        try:
            # Extract bullet points if available
            if any(marker in guidance for marker in ["•", "-", "*"]):
                lines = guidance.split("\n")
                bullets = [
                    line.strip()
                    for line in lines
                    if line.strip() and line.strip()[0] in ("•", "-", "*")
                ]
                if bullets:
                    return "\n".join(bullets[:8])

            # Truncate to 500 chars
            return guidance[:500] + "..." if len(guidance) > 500 else guidance

        except Exception as e:
            logger.error(f"❌ Error summarizing: {str(e)}")
            return guidance[:300]

    def _inject_context(self, template: str, context: Dict[str, Any]) -> str:
        """
        Inject context variables into prompt template.

        Replaces {{key}} with context values.

        Args:
            template: Prompt template with {{placeholders}}
            context: Dict of context values

        Returns:
            Prompt with injected context
        """
        prompt = template
        for key, value in context.items():
            placeholder = f"{{{{{key}}}}}"
            prompt = prompt.replace(placeholder, str(value))
        return prompt

    # ==================== RESPONSE PARSING ====================

    def _parse_response(self, response: str) -> Dict[str, Any]:
        """
        Parse AI response into structured document.

        Supports multiple formats:
        1. JSON with ```json``` wrapper
        2. Raw JSON object
        3. Plain text with markdown sections

        Args:
            response: Raw response from Claude

        Returns:
            Dict with 'generated_concept_document' and 'sections'
        """
        try:
            # Try JSON parsing first
            json_match = re.search(r"```json\s*(\{.*?\})\s*```", response, re.DOTALL)
            if json_match:
                parsed = json.loads(json_match.group(1))
                logger.info("📦 Parsed JSON from code block")
            else:
                # Try direct JSON
                json_match = re.search(r"\{.*?\}", response, re.DOTALL)
                if json_match:
                    parsed = json.loads(json_match.group(0))
                    logger.info("📦 Parsed JSON directly")
                else:
                    parsed = None

            if parsed:
                # Handle wrapped format
                if "concept_document" in parsed:
                    return {
                        "generated_concept_document": parsed["concept_document"].get(
                            "generated_concept_document", ""
                        ),
                        "sections": parsed["concept_document"].get("sections", {}),
                    }

                # Handle flat format
                if "generated_concept_document" in parsed:
                    return {
                        "generated_concept_document": parsed.get(
                            "generated_concept_document", ""
                        ),
                        "sections": parsed.get("sections", {}),
                    }

        except json.JSONDecodeError:
            logger.info("ℹ️  JSON parsing failed, using fallback")

        # Fallback to text extraction
        logger.info("📝 Extracting sections from text")
        return {
            "generated_concept_document": response,
            "sections": self._extract_sections_from_text(response),
        }

    def _extract_sections_from_text(self, text: str) -> Dict[str, str]:
        """
        Extract sections from markdown text.

        Looks for multiple section header formats:
        - # Section Title (markdown h1)
        - ## Section Title (markdown h2)
        - **Section Title** (bold format - common Claude output)

        Args:
            text: Markdown or plain text

        Returns:
            Dict of section_title: content
        """
        sections = {}
        current_section = None
        current_content = []

        # Regex pattern for **Bold Section Title** on its own line
        bold_pattern = re.compile(r"^\*\*([^*]+)\*\*\s*$")

        for line in text.split("\n"):
            stripped_line = line.strip()

            # Check for ## section headers
            if stripped_line.startswith("## "):
                if current_section:
                    sections[current_section] = "\n".join(current_content).strip()
                current_section = stripped_line[3:].strip()
                current_content = []
            # Check for # section headers
            elif stripped_line.startswith("# "):
                if current_section:
                    sections[current_section] = "\n".join(current_content).strip()
                current_section = stripped_line[2:].strip()
                current_content = []
            # Check for **Bold Title** format (common Claude output)
            elif bold_match := bold_pattern.match(stripped_line):
                if current_section:
                    sections[current_section] = "\n".join(current_content).strip()
                current_section = bold_match.group(1).strip()
                current_content = []
            elif current_section:
                current_content.append(line)

        if current_section:
            sections[current_section] = "\n".join(current_content).strip()

        logger.info(f"📊 Extracted {len(sections)} sections from text")
        for title in sections.keys():
            logger.info(f"   ✓ {title}")

        return sections


# ==================== SERVICE INSTANCE ====================

# Global instance for document generation
concept_generator = ConceptDocumentGenerator()
