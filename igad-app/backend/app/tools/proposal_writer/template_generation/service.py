"""Template Generation Service - Generates Word document from structure"""
import logging
from typing import Dict, Any, List, Optional
from datetime import datetime
from docx import Document
from io import BytesIO
from app.database.client import db_client

logger = logging.getLogger(__name__)
logger.setLevel(logging.INFO)


class TemplateGenerationService:
    def __init__(self):
        pass

    def generate_template(
        self,
        proposal_id: str,
        selected_sections: Optional[List[str]] = None,
        user_comments: Optional[Dict[str, str]] = None
    ) -> BytesIO:
        """
        Generate Word template from structure and workplan analysis.

        Args:
            proposal_id: Proposal code (PROP-YYYYMMDD-XXXX)
            selected_sections: List of section titles to include (None = all sections)
            user_comments: Dictionary of user comments per section {section_title: comment}

        Returns:
            BytesIO: Word document as bytes

        Raises:
            Exception: If proposal not found or structure missing
        """
        logger.info(f"📋 Generating template for proposal: {proposal_id}")
        logger.info(f"   Selected sections: {len(selected_sections) if selected_sections else 'all'}")
        logger.info(f"   User comments: {len(user_comments) if user_comments else 0}")

        # Load proposal
        proposal = db_client.get_item_sync(
            pk=f"PROPOSAL#{proposal_id}", sk="METADATA"
        )

        if not proposal:
            raise Exception(f"Proposal {proposal_id} not found")

        # Get structure workplan analysis
        structure_data = proposal.get("structure_workplan_analysis", {})
        if not structure_data:
            raise Exception("Structure and workplan analysis not found. Please complete Step 3 first.")

        # Unwrap if nested (structure_workplan_analysis.structure_workplan_analysis)
        structure_analysis = structure_data.get("structure_workplan_analysis", structure_data)

        # Get RFP analysis for context
        rfp_analysis = proposal.get("rfp_analysis", {})

        # Filter sections based on user selection
        mandatory_sections = structure_analysis.get("proposal_mandatory", [])
        outline_sections = structure_analysis.get("proposal_outline", [])

        # Combine all sections
        all_sections = mandatory_sections + outline_sections

        # Filter to selected sections if provided
        if selected_sections:
            filtered_sections = [
                s for s in all_sections
                if s.get("section_title") in selected_sections
            ]
            logger.info(f"   Filtered from {len(all_sections)} to {len(filtered_sections)} sections")
        else:
            filtered_sections = all_sections
            logger.info(f"   Using all {len(filtered_sections)} sections")

        # Mark which sections are mandatory
        mandatory_titles = {s.get("section_title") for s in mandatory_sections}
        for section in filtered_sections:
            section["is_mandatory"] = section.get("section_title") in mandatory_titles

        # Add user comments to sections
        if user_comments:
            for section in filtered_sections:
                title = section.get("section_title")
                if title and title in user_comments:
                    section["user_comment"] = user_comments[title]
                    logger.info(f"   ✓ Added comment for: {title}")
        
        # Create Word document
        doc = Document()

        # Add simple test content
        doc.add_heading('Proposal Template', 0)
        doc.add_paragraph('This is a test template document.')
        doc.add_paragraph('Generated successfully.')

        # Only add sections if they exist
        if filtered_sections and len(filtered_sections) > 0:
            doc.add_heading('Sections', 1)
            for section in filtered_sections[:3]:  # Only first 3 sections to test
                title = section.get("section_title", "Untitled")
                if title:
                    doc.add_heading(str(title), 2)
                    doc.add_paragraph("Section content placeholder")
        
        # Save to BytesIO
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        return buffer

    def _add_section(self, doc: Document, section: Dict[str, Any]):
        """
        Add a section to the document with all guidance and user comments.

        Args:
            doc: Document object to add section to
            section: Section data dictionary
        """
        # Section title
        section_title = str(section.get("section_title", "Untitled Section"))
        doc.add_heading(section_title, 2)

        # Mandatory badge
        if section.get("is_mandatory"):
            doc.add_paragraph("MANDATORY SECTION")

        # Word count
        word_count = section.get("recommended_word_count")
        if word_count:
            doc.add_paragraph("Recommended length: " + str(word_count))

        doc.add_paragraph()  # Spacing

        # Purpose
        purpose = section.get("purpose")
        if purpose:
            doc.add_heading("Purpose", 3)
            doc.add_paragraph(str(purpose))

        # User comment (if provided)
        user_comment = section.get("user_comment")
        if user_comment:
            doc.add_heading("Your Notes & Context", 3)
            doc.add_paragraph(str(user_comment))

        # Content guidance
        guidance = section.get("content_guidance")
        if guidance:
            doc.add_heading("What to Include", 3)
            doc.add_paragraph(str(guidance))

        # Guiding questions
        questions = section.get("guiding_questions", [])
        if questions and isinstance(questions, list):
            doc.add_heading("Guiding Questions", 3)
            for q in questions:
                if q:
                    doc.add_paragraph(str(q))

        # Writing space
        doc.add_paragraph()
        doc.add_heading("Your Content", 3)
        doc.add_paragraph("[Write your content here]")

        # Section separator
        doc.add_paragraph()
        doc.add_paragraph("=" * 50)
        doc.add_paragraph()
