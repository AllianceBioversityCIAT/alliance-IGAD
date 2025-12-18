"""
Proposal Template Generation Service

Generates full draft proposals based on:
- Structure workplan analysis (selected sections)
- Concept document (improved concept)
- RFP analysis and requirements
- Reference proposals analysis
- Existing work analysis

The service uses Claude to generate a complete, donor-aligned draft proposal
document following the structure and guidance from previous steps.
"""

import json
import logging
import os
import re
import traceback
from datetime import datetime
from decimal import Decimal
from io import BytesIO
from typing import Any, Dict, List, Optional

import boto3
from boto3.dynamodb.conditions import Attr
from docx import Document

from app.shared.ai.bedrock_service import BedrockService
from app.tools.proposal_writer.proposal_template_generation.config import (
    PROPOSAL_TEMPLATE_GENERATION_SETTINGS,
)

logger = logging.getLogger(__name__)
logger.setLevel(logging.INFO)


class ProposalTemplateGenerator:
    """
    Generates full draft proposals using AI.

    Workflow:
    1. Load prompt template from DynamoDB (Prompt 4.5 - Draft Proposal)
    2. Prepare proposal structure from selected sections
    3. Prepare context with all analysis data
    4. Inject context into prompt
    5. Call Claude via Bedrock
    6. Parse and structure the generated document
    7. Optionally save to S3
    """

    def __init__(self):
        """
        Initialize Bedrock, DynamoDB, and S3 clients.

        Sets up clients for:
        - Bedrock: Claude AI invocation
        - DynamoDB: Proposal data and prompt retrieval
        - S3: Document storage
        """
        self.bedrock = BedrockService()
        self.dynamodb = boto3.resource("dynamodb")
        self.s3 = boto3.client("s3")
        self.table_name = os.environ.get("DYNAMODB_TABLE", "igad-testing-main-table")
        self.bucket = os.environ.get("PROPOSALS_BUCKET")
        if not self.bucket:
            logger.warning("⚠️ PROPOSALS_BUCKET environment variable not set")

    def generate_template(
        self,
        proposal_code: str,
        selected_sections: List[str],
        rfp_analysis: Dict[str, Any],
        concept_document: Dict[str, Any],
        structure_workplan_analysis: Dict[str, Any],
        reference_proposals_analysis: Optional[Dict[str, Any]] = None,
        existing_work_analysis: Optional[Dict[str, Any]] = None,
        user_comments: Optional[Dict[str, str]] = None,
    ) -> Dict[str, Any]:
        """
        Generate full draft proposal using AI.

        Args:
            proposal_code: Proposal identifier for logging
            selected_sections: List of section titles selected by user
            rfp_analysis: RFP analysis from step 1
            concept_document: Improved concept document from step 2
            structure_workplan_analysis: Structure workplan from step 3
            reference_proposals_analysis: Optional reference proposals analysis
            existing_work_analysis: Optional existing work analysis
            user_comments: Optional user comments per section

        Returns:
            Dict with 'generated_proposal', 'sections', and metadata

        Raises:
            ValueError: If prompt template not found
            Exception: If document generation fails
        """
        try:
            logger.info(f"📋 Generating proposal template for: {proposal_code}")
            logger.info(f"   Selected sections: {len(selected_sections)}")

            # Step 1: Load prompt template
            logger.info("📝 Loading prompt template (Prompt 4.5 - Draft Proposal)...")
            prompt_parts = self._get_prompt_template()

            if not prompt_parts:
                raise ValueError(
                    "Prompt template not found in DynamoDB. "
                    "Ensure Prompt 4.5 (Draft Proposal) is configured."
                )

            # Step 2: Prepare proposal structure
            logger.info("🔄 Preparing proposal structure...")
            proposal_structure = self._prepare_proposal_structure(
                structure_workplan_analysis,
                selected_sections,
                user_comments,
            )

            # Step 3: Prepare concept document text
            logger.info("📄 Preparing concept document...")
            concept_text = self._prepare_concept_document(concept_document)

            # Step 4: Prepare context
            logger.info("🔄 Preparing context...")
            context = self._prepare_context(
                proposal_structure=proposal_structure,
                concept_text=concept_text,
                rfp_analysis=rfp_analysis,
                reference_proposals_analysis=reference_proposals_analysis,
                existing_work_analysis=existing_work_analysis,
            )

            # Step 5: Build final prompt
            user_prompt = self._inject_context(prompt_parts["user_prompt"], context)
            final_prompt = f"{user_prompt}\n\n{prompt_parts['output_format']}".strip()

            # Step 6: Call Bedrock
            logger.info("📡 Calling Bedrock (this may take 5-10 minutes)...")
            start_time = datetime.utcnow()

            ai_response = self.bedrock.invoke_claude(
                system_prompt=prompt_parts["system_prompt"],
                user_prompt=final_prompt,
                max_tokens=PROPOSAL_TEMPLATE_GENERATION_SETTINGS.get(
                    "max_tokens", 32000
                ),
                temperature=PROPOSAL_TEMPLATE_GENERATION_SETTINGS.get(
                    "temperature", 0.2
                ),
                model_id=PROPOSAL_TEMPLATE_GENERATION_SETTINGS["model"],
            )

            elapsed = (datetime.utcnow() - start_time).total_seconds()
            logger.info(f"✅ Response received in {elapsed:.1f} seconds")

            # Step 7: Parse response
            logger.info("📊 Parsing response...")
            document = self._parse_response(ai_response)

            # Add metadata (using Decimal for float to avoid DynamoDB Float error)
            document["metadata"] = {
                "proposal_code": proposal_code,
                "generated_at": datetime.utcnow().isoformat(),
                "generation_time_seconds": Decimal(str(round(elapsed, 2))),
                "sections_count": len(selected_sections),
                "selected_sections": selected_sections,
            }

            logger.info("✅ Proposal template generated successfully")
            return document

        except Exception as e:
            logger.error(f"❌ Template generation failed: {str(e)}")
            traceback.print_exc()
            raise

    # ==================== PROMPT AND CONTEXT ====================

    def _get_prompt_template(self) -> Optional[Dict[str, str]]:
        """
        Load template generation prompt from DynamoDB.

        Filters for:
        - is_active: true
        - section: "proposal_writer"
        - sub_section: "step-3"
        - categories: contains "Draft Proposal"

        Returns:
            Dict with 'system_prompt', 'user_prompt', 'output_format', or None
        """
        try:
            table = self.dynamodb.Table(self.table_name)
            filter_expr = (
                Attr("is_active").eq(True)
                & Attr("section").eq(
                    PROPOSAL_TEMPLATE_GENERATION_SETTINGS["section"]
                )
                & Attr("sub_section").eq(
                    PROPOSAL_TEMPLATE_GENERATION_SETTINGS["sub_section"]
                )
                & Attr("categories").contains(
                    PROPOSAL_TEMPLATE_GENERATION_SETTINGS["category"]
                )
            )

            # Handle DynamoDB pagination
            items = []
            response = table.scan(FilterExpression=filter_expr)
            items.extend(response.get("Items", []))

            while "LastEvaluatedKey" in response:
                response = table.scan(
                    FilterExpression=filter_expr,
                    ExclusiveStartKey=response["LastEvaluatedKey"]
                )
                items.extend(response.get("Items", []))

            if not items:
                logger.warning("⚠️ No prompts found in DynamoDB for Draft Proposal")
                return None

            prompt_item = items[0]
            logger.info(f"✅ Loaded prompt: {prompt_item.get('name', 'Unnamed')}")

            return {
                "system_prompt": prompt_item.get("system_prompt", ""),
                "user_prompt": prompt_item.get("user_prompt_template", ""),
                "output_format": prompt_item.get("output_format", ""),
            }

        except Exception as e:
            logger.error(f"❌ Error loading prompt: {str(e)}")
            return None

    def _prepare_proposal_structure(
        self,
        structure_workplan_analysis: Dict[str, Any],
        selected_sections: List[str],
        user_comments: Optional[Dict[str, str]] = None,
    ) -> Dict[str, Any]:
        """
        Prepare proposal structure from selected sections.

        Filters mandatory and outline sections to only include selected ones.
        Adds user comments if provided.

        Args:
            structure_workplan_analysis: Full structure workplan data
            selected_sections: List of section titles to include
            user_comments: Optional user comments per section

        Returns:
            Filtered proposal structure with selected sections
        """
        try:
            # Unwrap nested structure if present
            analysis = structure_workplan_analysis.get(
                "structure_workplan_analysis", structure_workplan_analysis
            )

            mandatory_sections = analysis.get("proposal_mandatory", [])
            outline_sections = analysis.get("proposal_outline", [])

            logger.info(
                f"📊 Total sections: {len(mandatory_sections)} mandatory + "
                f"{len(outline_sections)} outline"
            )

            # Filter sections based on selection
            selected_set = set(selected_sections)

            filtered_mandatory = [
                s for s in mandatory_sections
                if s.get("section_title") in selected_set
            ]
            filtered_outline = [
                s for s in outline_sections
                if s.get("section_title") in selected_set
            ]

            # Add user comments if provided
            if user_comments:
                for section in filtered_mandatory + filtered_outline:
                    title = section.get("section_title")
                    if title and title in user_comments:
                        section["user_comment"] = user_comments[title]
                        logger.info(f"   📝 Added comment for: {title}")

            logger.info(
                f"✅ Filtered to {len(filtered_mandatory)} mandatory + "
                f"{len(filtered_outline)} outline sections"
            )

            return {
                "proposal_mandatory": filtered_mandatory,
                "proposal_outline": filtered_outline,
            }

        except Exception as e:
            logger.error(f"❌ Error preparing proposal structure: {str(e)}")
            return {"proposal_mandatory": [], "proposal_outline": []}

    def _prepare_concept_document(
        self,
        concept_document: Dict[str, Any],
    ) -> str:
        """
        Extract and prepare concept document text.

        Args:
            concept_document: Concept document data from step 2

        Returns:
            Concept document as text string
        """
        try:
            # Handle nested structure
            doc = concept_document.get("concept_document_v2", concept_document)
            if isinstance(doc, dict):
                doc = doc.get("concept_document_v2", doc)

            # Extract generated document text
            if isinstance(doc, dict):
                generated_text = doc.get("generated_concept_document", "")
                if generated_text:
                    logger.info(
                        f"✅ Extracted concept document: {len(generated_text)} chars"
                    )
                    return self._truncate_text(generated_text, 50000)

                # Fallback: concatenate sections
                sections = doc.get("sections", {})
                if sections:
                    text = "\n\n".join(
                        f"## {title}\n{content}"
                        for title, content in sections.items()
                    )
                    logger.info(f"✅ Built from sections: {len(text)} chars")
                    return self._truncate_text(text, 50000)

            # If it's already a string
            if isinstance(doc, str):
                return self._truncate_text(doc, 50000)

            logger.warning("⚠️ Could not extract concept document text")
            return ""

        except Exception as e:
            logger.error(f"❌ Error preparing concept document: {str(e)}")
            return ""

    def _prepare_context(
        self,
        proposal_structure: Dict[str, Any],
        concept_text: str,
        rfp_analysis: Dict[str, Any],
        reference_proposals_analysis: Optional[Dict[str, Any]] = None,
        existing_work_analysis: Optional[Dict[str, Any]] = None,
    ) -> Dict[str, str]:
        """
        Prepare context for prompt injection.

        Args:
            proposal_structure: Filtered proposal sections
            concept_text: Concept document text
            rfp_analysis: RFP analysis data
            reference_proposals_analysis: Optional reference analysis
            existing_work_analysis: Optional existing work analysis

        Returns:
            Dict with context strings for prompt injection
        """
        context = {
            "PROPOSAL STRUCTURE": json.dumps(proposal_structure, indent=2),
            "INITIAL CONCEPT": concept_text,
            "RFP ANALYSIS": json.dumps(rfp_analysis, indent=2),
        }

        # Add reference proposals analysis if provided
        if reference_proposals_analysis:
            # Unwrap nested structure if present
            analysis = reference_proposals_analysis.get(
                "reference_proposals_analysis", reference_proposals_analysis
            )
            context["REFERENCE PROPOSALS ANALYSIS"] = json.dumps(analysis, indent=2)
            logger.info("✅ Reference proposals analysis added to context")
        else:
            context["REFERENCE PROPOSALS ANALYSIS"] = "Not provided"

        # Add existing work analysis if provided
        if existing_work_analysis:
            # Unwrap nested structure if present
            analysis = existing_work_analysis.get(
                "existing_work_analysis", existing_work_analysis
            )
            context["EXISTING WORK ANALYSIS"] = json.dumps(analysis, indent=2)
            logger.info("✅ Existing work analysis added to context")
        else:
            context["EXISTING WORK ANALYSIS"] = "Not provided"

        return context

    def _inject_context(self, template: str, context: Dict[str, str]) -> str:
        """
        Inject context variables into prompt template.

        Replaces {[KEY]} with context values.
        Note: Uses {[KEY]} format (different from {{key}} in other services).

        Args:
            template: Prompt template with {[PLACEHOLDER]} markers
            context: Dict of context values

        Returns:
            Prompt with injected context
        """
        prompt = template
        for key, value in context.items():
            # Format: {[KEY]}
            placeholder = "{[" + key + "]}"
            prompt = prompt.replace(placeholder, str(value))
        return prompt

    # ==================== RESPONSE PARSING ====================

    def _parse_response(self, response: str) -> Dict[str, Any]:
        """
        Parse AI response into structured document.

        The response is expected to be a continuous markdown document
        with section headers.

        Args:
            response: Raw response from Claude

        Returns:
            Dict with 'generated_proposal' and 'sections'
        """
        try:
            # Extract sections from markdown
            sections = self._extract_sections_from_text(response)

            return {
                "generated_proposal": response,
                "sections": sections,
            }

        except Exception as e:
            logger.error(f"❌ Error parsing response: {str(e)}")
            return {
                "generated_proposal": response,
                "sections": {},
            }

    def _extract_sections_from_text(self, text: str) -> Dict[str, str]:
        """
        Extract sections from markdown text.

        Looks for multiple section header formats:
        - # Section Title (markdown h1)
        - ## Section Title (markdown h2)
        - **Section Title** (bold format - common Claude output)

        Args:
            text: Markdown or plain text

        Returns:
            Dict of section_title: content
        """
        sections = {}
        current_section = None
        current_content = []

        # Regex pattern for **Bold Section Title** on its own line
        bold_pattern = re.compile(r"^\*\*([^*]+)\*\*\s*$")

        for line in text.split("\n"):
            stripped_line = line.strip()

            # Check for ## section headers
            if stripped_line.startswith("## "):
                if current_section:
                    sections[current_section] = "\n".join(current_content).strip()
                current_section = stripped_line[3:].strip()
                current_content = []
            # Check for # section headers
            elif stripped_line.startswith("# "):
                if current_section:
                    sections[current_section] = "\n".join(current_content).strip()
                current_section = stripped_line[2:].strip()
                current_content = []
            # Check for **Bold Title** format (common Claude output)
            elif bold_match := bold_pattern.match(stripped_line):
                if current_section:
                    sections[current_section] = "\n".join(current_content).strip()
                current_section = bold_match.group(1).strip()
                current_content = []
            elif current_section:
                current_content.append(line)

        if current_section:
            sections[current_section] = "\n".join(current_content).strip()

        logger.info(f"📊 Extracted {len(sections)} sections from text")
        for title in list(sections.keys())[:5]:  # Log first 5
            logger.info(f"   ✓ {title}")
        if len(sections) > 5:
            logger.info(f"   ... and {len(sections) - 5} more sections")

        return sections

    # ==================== UTILITY METHODS ====================

    def _truncate_text(self, text: str, max_chars: int) -> str:
        """
        Intelligently truncate text if needed.

        Keeps 60% from beginning and 40% from end.

        Args:
            text: Text to truncate
            max_chars: Maximum characters

        Returns:
            Truncated text with notice if truncated
        """
        if len(text) <= max_chars:
            return text

        logger.info(f"✂️ Truncating text from {len(text)} to {max_chars} chars")

        chars_to_keep = max_chars - 200
        beginning_chars = int(chars_to_keep * 0.6)
        ending_chars = int(chars_to_keep * 0.4)

        beginning = text[:beginning_chars]
        ending = text[-ending_chars:]

        return (
            f"{beginning}\n\n"
            f"[... Middle section truncated - "
            f"total document: {len(text)} characters ...]\n\n"
            f"{ending}"
        )

    def save_to_s3(
        self,
        proposal_code: str,
        content: str,
        filename: str = "draft_proposal.md",
    ) -> Optional[str]:
        """
        Save generated document to S3.

        Args:
            proposal_code: Proposal code for path
            content: Document content
            filename: Filename to save as

        Returns:
            S3 URL if successful, None otherwise
        """
        if not self.bucket:
            logger.warning("⚠️ Cannot save to S3: bucket not configured")
            return None

        try:
            s3_key = f"{proposal_code}/documents/proposal_template/{filename}"

            self.s3.put_object(
                Bucket=self.bucket,
                Key=s3_key,
                Body=content.encode("utf-8"),
                ContentType="text/markdown",
            )

            s3_url = f"s3://{self.bucket}/{s3_key}"
            logger.info(f"✅ Saved to S3: {s3_url}")
            return s3_url

        except Exception as e:
            logger.error(f"❌ Error saving to S3: {str(e)}")
            return None

    def generate_docx(self, content: str, sections: Dict[str, str]) -> BytesIO:
        """
        Generate Word document from content.

        Args:
            content: Full document content
            sections: Dict of section_title: content

        Returns:
            BytesIO buffer with DOCX content
        """
        try:
            doc = Document()

            # Title
            doc.add_heading("Draft Proposal", 0)
            doc.add_paragraph(
                f"Generated on: {datetime.utcnow().strftime('%Y-%m-%d %H:%M UTC')}"
            )
            doc.add_paragraph()

            # Add sections
            for title, section_content in sections.items():
                doc.add_heading(title, 1)
                # Add paragraphs
                for paragraph in section_content.split("\n\n"):
                    if paragraph.strip():
                        doc.add_paragraph(paragraph.strip())
                doc.add_paragraph()  # Spacing between sections

            # Save to buffer
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)

            logger.info(f"✅ Generated DOCX with {len(sections)} sections")
            return buffer

        except Exception as e:
            logger.error(f"❌ Error generating DOCX: {str(e)}")
            raise


# ==================== SERVICE INSTANCE ====================

# Global instance for template generation
proposal_template_generator = ProposalTemplateGenerator()
