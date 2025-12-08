"""Vector Embeddings Service for S3 Vectors"""
import boto3
import json
from typing import List, Dict, Any, Optional
from datetime import datetime

class VectorEmbeddingsService:
    def __init__(self):
        self.bedrock = boto3.client('bedrock-runtime', region_name='us-east-1')
        self.s3vectors = boto3.client('s3vectors', region_name='us-east-1')
        self.s3 = boto3.client('s3', region_name='us-east-1')
        self.bucket_name = "igad-proposals-vectors-testing"
        self.documents_bucket = "igad-proposal-documents-569113802249"
        self.model_id = "amazon.titan-embed-text-v2:0"
    
    def generate_embedding(self, text: str) -> List[float]:
        """Generate embedding using Bedrock Titan"""
        response = self.bedrock.invoke_model(
            modelId=self.model_id,
            body=json.dumps({"inputText": text})
        )
        return json.loads(response['body'].read())['embedding']
    
    def insert_reference_proposal(self, proposal_id: str, text: str, metadata: Dict[str, str]) -> bool:
        """Insert reference proposal vector with metadata encoded in key"""
        try:
            embedding = self.generate_embedding(text)
            
            # Encode metadata in key (bypass AWS metadata bug)
            key = "|".join([
                proposal_id,
                metadata.get("donor", ""),
                metadata.get("sector", ""),
                metadata.get("year", ""),
                metadata.get("document_name", ""),
                metadata.get("chunk_index", "0"),
                metadata.get("total_chunks", "1")
            ])
            
            self.s3vectors.put_vectors(
                vectorBucketName=self.bucket_name,
                indexName="reference-proposals-index",
                vectors=[{
                    "key": key,
                    "data": {"float32": embedding},
                    "metadata": {}  # Empty - using key encoding instead
                }]
            )
            return True
        except Exception as e:
            print(f"Error: {e}")
            return False
    
    def insert_existing_work(self, proposal_id: str, text: str, metadata: Dict[str, str]) -> bool:
        """
        Insert existing work vector with metadata encoded in key.

        Key format: proposal_id|organization|project_type|region|document_name|chunk_index|total_chunks
        This matches the reference proposals format for consistency.
        """
        try:
            embedding = self.generate_embedding(text)

            # Encode metadata in key (bypass AWS metadata bug)
            # Format: proposal_id|organization|project_type|region|document_name|chunk_index|total_chunks
            key = "|".join([
                proposal_id,
                metadata.get("organization", ""),
                metadata.get("project_type", ""),
                metadata.get("region", ""),
                metadata.get("document_name", ""),
                metadata.get("chunk_index", "0"),
                metadata.get("total_chunks", "1")
            ])
            
            self.s3vectors.put_vectors(
                vectorBucketName=self.bucket_name,
                indexName="existing-work-index",
                vectors=[{
                    "key": key,
                    "data": {"float32": embedding},
                    "metadata": {}  # Empty - using key encoding instead
                }]
            )
            return True
        except Exception as e:
            print(f"Error: {e}")
            return False
    
    def search_similar_proposals(self, query_text: str, top_k: int = 5, 
                                 filters: Optional[Dict[str, str]] = None) -> List[Dict[str, Any]]:
        """Search similar reference proposals"""
        try:
            query_embedding = self.generate_embedding(query_text)
            params = {
                "vectorBucketName": self.bucket_name,
                "indexName": "reference-proposals-index",
                "queryVector": {"float32": query_embedding},
                "topK": top_k,
                "returnDistance": True,
                "returnMetadata": True
            }
            if filters:
                params["filter"] = filters
            return self.s3vectors.query_vectors(**params).get('vectors', [])
        except Exception as e:
            print(f"Error: {e}")
            return []
    
    def search_similar_work(self, query_text: str, top_k: int = 5,
                           filters: Optional[Dict[str, str]] = None) -> List[Dict[str, Any]]:
        """Search similar existing work"""
        try:
            query_embedding = self.generate_embedding(query_text)
            params = {
                "vectorBucketName": self.bucket_name,
                "indexName": "existing-work-index",
                "queryVector": {"float32": query_embedding},
                "topK": top_k,
                "returnDistance": True,
                "returnMetadata": True
            }
            if filters:
                params["filter"] = filters
            return self.s3vectors.query_vectors(**params).get('vectors', [])
        except Exception as e:
            print(f"Error: {e}")
            return []
    
    def delete_proposal_vectors(self, proposal_id: str) -> bool:
        """
        Delete all vectors for a proposal from both indices (using key encoding).

        Args:
            proposal_id: Proposal code to delete vectors for

        Returns:
            True if deletion successful, False otherwise
        """
        try:
            total_deleted = 0

            for index in ["reference-proposals-index", "existing-work-index"]:
                print(f"   Checking {index}...")
                response = self.s3vectors.list_vectors(
                    vectorBucketName=self.bucket_name,
                    indexName=index
                )

                # Parse keys and filter by proposal_id
                keys = []
                for v in response.get('vectors', []):
                    key = v.get('key', '')
                    parts = key.split('|')
                    if len(parts) >= 1 and parts[0] == proposal_id:
                        keys.append(key)

                if keys:
                    self.s3vectors.delete_vectors(
                        vectorBucketName=self.bucket_name,
                        indexName=index,
                        keys=keys
                    )
                    print(f"   ✅ Deleted {len(keys)} vectors from {index}")
                    total_deleted += len(keys)
                else:
                    print(f"   ℹ️  No vectors found in {index}")

            if total_deleted > 0:
                print(f"✅ Total vectors deleted: {total_deleted}")
                return True
            else:
                print(f"ℹ️  No vectors found for proposal {proposal_id}")
                return True  # Not an error, just nothing to delete

        except Exception as e:
            print(f"❌ Error deleting vectors: {e}")
            import traceback
            traceback.print_exc()
            return False

    def delete_vectors_by_document_name(self, document_name: str, index_name: str) -> bool:
        """
        Delete all vectors associated with a specific document name (using key encoding)

        Args:
            document_name: The document filename to delete vectors for
            index_name: The index to delete from

        Returns:
            True if successful, False otherwise
        """
        try:
            response = self.s3vectors.list_vectors(
                vectorBucketName=self.bucket_name,
                indexName=index_name
            )

            # Parse keys and filter by document_name
            keys_to_delete = []
            for vector in response.get('vectors', []):
                key = vector.get('key', '')
                parts = key.split('|')
                if len(parts) >= 7 and parts[4] == document_name:
                    keys_to_delete.append(key)

            if keys_to_delete:
                print(f"Deleting {len(keys_to_delete)} vectors for document {document_name}")
                self.s3vectors.delete_vectors(
                    vectorBucketName=self.bucket_name,
                    indexName=index_name,
                    keys=keys_to_delete
                )
                return True
            else:
                print(f"No vectors found for document {document_name}")
                return True

        except Exception as e:
            print(f"Error deleting vectors by document name: {e}")
            import traceback
            traceback.print_exc()
            return False
    
    def get_documents_by_proposal(self, proposal_id: str, index_name: str = "reference-proposals-index", max_docs: int = 5) -> List[Dict[str, Any]]:
        """
        Retrieve documents for a proposal from S3 Vectors (using key encoding)
        
        Args:
            proposal_id: Proposal ID to filter by
            index_name: Index to search in
            max_docs: Maximum number of documents to return
            
        Returns:
            List of documents with their full text reconstructed from S3
        """
        try:
            response = self.s3vectors.list_vectors(
                vectorBucketName=self.bucket_name,
                indexName=index_name
            )
            
            vectors = response.get('vectors', [])
            
            # Parse keys and filter by proposal_id
            proposal_vectors = []
            for vector in vectors:
                key = vector.get('key', '')
                parts = key.split('|')
                if len(parts) >= 7 and parts[0] == proposal_id:
                    vector['decoded_metadata'] = {
                        'proposal_id': parts[0],
                        'donor': parts[1],
                        'sector': parts[2],
                        'year': parts[3],
                        'document_name': parts[4],
                        'chunk_index': int(parts[5]),
                        'total_chunks': parts[6]
                    }
                    proposal_vectors.append(vector)
            
            if not proposal_vectors:
                return []
            
            # Group by document_name
            docs_by_name = {}
            for vector in proposal_vectors:
                metadata = vector['decoded_metadata']
                doc_name = metadata['document_name']
                
                if doc_name not in docs_by_name:
                    docs_by_name[doc_name] = []
                
                docs_by_name[doc_name].append(vector)
            
            # Reconstruct each document with full text from S3
            reconstructed_docs = []
            for doc_name, doc_vectors in docs_by_name.items():
                doc_vectors.sort(key=lambda v: v['decoded_metadata']['chunk_index'])
                first_metadata = doc_vectors[0]['decoded_metadata']

                # Retrieve full text from S3
                try:
                    # Correct S3 key path for reference proposals
                    s3_key = f"{proposal_id}/documents/references/{doc_name}"
                    s3_response = self.s3.get_object(
                        Bucket=self.documents_bucket,
                        Key=s3_key
                    )

                    # Extract text based on file format
                    content = s3_response['Body'].read()
                    ext = doc_name.lower().split('.')[-1] if '.' in doc_name else ''

                    if ext == 'pdf':
                        # PDF extraction
                        from PyPDF2 import PdfReader
                        from io import BytesIO
                        pdf_reader = PdfReader(BytesIO(content))
                        full_text = "\n\n".join([page.extract_text() for page in pdf_reader.pages])

                    elif ext == 'docx':
                        # DOCX extraction
                        from docx import Document
                        from io import BytesIO
                        doc = Document(BytesIO(content))
                        full_text = "\n\n".join([para.text for para in doc.paragraphs if para.text.strip()])

                    elif ext == 'txt':
                        # TXT extraction
                        full_text = content.decode('utf-8')

                    else:
                        # Unsupported format
                        print(f"⚠️  Unsupported format: {ext} for {doc_name}")
                        full_text = f"[Unsupported format: {ext}]"

                    reconstructed_docs.append({
                        "document_name": doc_name,
                        "full_text": full_text.strip(),
                        "chunk_count": len(doc_vectors),
                        "metadata": first_metadata
                    })

                    print(f"✅ Reconstructed {doc_name} ({ext.upper()}, {len(doc_vectors)} chunks)")

                except Exception as e:
                    print(f"❌ Error reading {doc_name} from S3: {e}")
                    import traceback
                    traceback.print_exc()
                    reconstructed_docs.append({
                        "document_name": doc_name,
                        "full_text": f"[Error retrieving text: {str(e)}]",
                        "chunk_count": len(doc_vectors),
                        "metadata": first_metadata
                    })
            
            return reconstructed_docs[:max_docs]
            
        except Exception as e:
            print(f"Error retrieving documents by proposal: {e}")
            import traceback
            traceback.print_exc()
            return []

    def search_and_reconstruct_proposals(
        self,
        query_text: str,
        top_k: int = 5,
        index_name: str = "reference-proposals-index"
    ) -> List[Dict[str, Any]]:
        """
        Search similar proposals using semantic query and reconstruct full documents.

        This method:
        1. Generates embedding from semantic query
        2. Searches vector index for similar chunks
        3. Groups chunks by document
        4. Ranks documents by average similarity
        5. Reconstructs top-K full documents from S3

        Args:
            query_text: Semantic query from RFP analysis
            top_k: Number of top documents to return
            index_name: Vector index to search (default: reference-proposals-index)

        Returns:
            List of reconstructed documents with full text, sorted by relevance
            Each document contains:
            - document_name: Filename
            - full_text: Complete document text
            - similarity_score: 0-1 score (1 = most similar)
            - chunks_matched: Number of chunks that matched
            - metadata: Decoded metadata from key
        """
        try:
            print(f"🔍 Generating embedding for semantic search...")
            # 1. Generate query embedding
            query_embedding = self.generate_embedding(query_text)

            # 2. Semantic search (get more chunks to ensure good document coverage)
            print(f"🔎 Searching {index_name} for top {top_k * 10} relevant chunks...")
            response = self.s3vectors.query_vectors(
                vectorBucketName=self.bucket_name,
                indexName=index_name,
                queryVector={"float32": query_embedding},
                topK=top_k * 10,  # Get 10x chunks to group into top_k documents
                returnDistance=True,
                returnMetadata=True
            )

            vectors = response.get('vectors', [])

            if not vectors:
                print("⚠️  No similar vectors found")
                return []

            print(f"✅ Found {len(vectors)} matching chunks")

            # 3. Group chunks by document
            docs_by_name = {}
            for vector in vectors:
                key = vector.get('key', '')
                parts = key.split('|')

                # Key format: proposal_id|donor|sector|year|document_name|chunk_index|total_chunks
                if len(parts) >= 5:
                    doc_name = parts[4]  # document_name

                    if doc_name not in docs_by_name:
                        docs_by_name[doc_name] = {
                            'chunks': [],
                            'avg_distance': 0,
                            'metadata': {
                                'proposal_id': parts[0],
                                'donor': parts[1],
                                'sector': parts[2],
                                'year': parts[3],
                                'document_name': doc_name
                            }
                        }

                    docs_by_name[doc_name]['chunks'].append({
                        'key': key,
                        'distance': vector.get('distance', 1.0),
                        'chunk_index': int(parts[5]) if len(parts) > 5 else 0
                    })

            print(f"📊 Grouped into {len(docs_by_name)} unique documents")

            # 4. Calculate avg_distance per document (lower = more similar)
            for doc in docs_by_name.values():
                distances = [c['distance'] for c in doc['chunks']]
                doc['avg_distance'] = sum(distances) / len(distances)

            # 5. Sort by relevance (lowest distance = highest similarity)
            sorted_docs = sorted(
                docs_by_name.items(),
                key=lambda x: x[1]['avg_distance']
            )[:top_k]

            print(f"🎯 Selected top {len(sorted_docs)} most relevant documents")

            # 6. Reconstruct full documents from S3
            reconstructed = []
            for doc_name, doc_info in sorted_docs:
                try:
                    # Use first chunk metadata to find S3 path
                    proposal_id = doc_info['metadata']['proposal_id']

                    # S3 path depends on index type
                    if index_name == "existing-work-index":
                        s3_key = f"{proposal_id}/documents/supporting/{doc_name}"
                    else:  # reference-proposals-index
                        s3_key = f"{proposal_id}/documents/references/{doc_name}"

                    print(f"   📄 Reconstructing: {doc_name} (similarity: {1 - doc_info['avg_distance']:.2%})")

                    s3_response = self.s3.get_object(
                        Bucket=self.documents_bucket,
                        Key=s3_key
                    )

                    # Extract text based on file format
                    content = s3_response['Body'].read()
                    ext = doc_name.lower().split('.')[-1] if '.' in doc_name else ''

                    if ext == 'pdf':
                        from PyPDF2 import PdfReader
                        from io import BytesIO
                        pdf_reader = PdfReader(BytesIO(content))
                        full_text = "\n\n".join([page.extract_text() for page in pdf_reader.pages])
                    elif ext == 'docx':
                        from docx import Document
                        from io import BytesIO
                        doc = Document(BytesIO(content))
                        full_text = "\n\n".join([para.text for para in doc.paragraphs if para.text.strip()])
                    elif ext == 'txt':
                        full_text = content.decode('utf-8')
                    else:
                        print(f"      ⚠️  Unsupported format: {ext}")
                        full_text = f"[Unsupported format: {ext}]"

                    reconstructed.append({
                        "document_name": doc_name,
                        "full_text": full_text.strip(),
                        "similarity_score": 1 - doc_info['avg_distance'],  # Convert distance to similarity
                        "chunks_matched": len(doc_info['chunks']),
                        "metadata": doc_info['metadata']
                    })

                    print(f"      ✅ Reconstructed ({len(full_text)} chars, {len(doc_info['chunks'])} chunks)")

                except Exception as e:
                    print(f"      ❌ Error reconstructing {doc_name}: {e}")
                    import traceback
                    traceback.print_exc()
                    continue

            print(f"✅ Successfully reconstructed {len(reconstructed)} documents")
            return reconstructed

        except Exception as e:
            print(f"❌ Error in semantic search: {e}")
            import traceback
            traceback.print_exc()
            return []

