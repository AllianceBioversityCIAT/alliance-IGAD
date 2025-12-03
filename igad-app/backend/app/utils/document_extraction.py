"""
Document text extraction utilities
Supports PDF and DOCX files
"""
from io import BytesIO
from typing import Optional
import PyPDF2
from docx import Document


def extract_text_from_pdf(file_bytes: bytes) -> Optional[str]:
    """
    Extract text from PDF file bytes using PyPDF2

    Args:
        file_bytes: PDF file content as bytes

    Returns:
        Extracted text or None if extraction fails
    """
    try:
        pdf_file = BytesIO(file_bytes)
        pdf_reader = PyPDF2.PdfReader(pdf_file)

        text_content = []
        for page in pdf_reader.pages:
            page_text = page.extract_text()
            if page_text:
                text_content.append(page_text)

        full_text = "\n\n".join(text_content)

        # Return None if no text was extracted
        if not full_text.strip():
            return None

        return full_text

    except Exception as e:
        print(f"PDF extraction error: {str(e)}")
        return None


def extract_text_from_docx(file_bytes: bytes) -> Optional[str]:
    """
    Extract text from DOCX file bytes using python-docx

    Args:
        file_bytes: DOCX file content as bytes

    Returns:
        Extracted text or None if extraction fails
    """
    try:
        docx_file = BytesIO(file_bytes)
        doc = Document(docx_file)

        text_content = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_content.append(paragraph.text)

        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_content.append(cell.text)

        full_text = "\n\n".join(text_content)

        # Return None if no text was extracted
        if not full_text.strip():
            return None

        return full_text

    except Exception as e:
        print(f"DOCX extraction error: {str(e)}")
        return None


def extract_text_from_file(file_bytes: bytes, filename: str) -> Optional[str]:
    """
    Extract text from file based on extension

    Args:
        file_bytes: File content as bytes
        filename: Original filename to determine type

    Returns:
        Extracted text or None if extraction fails
    """
    filename_lower = filename.lower()

    if filename_lower.endswith('.pdf'):
        return extract_text_from_pdf(file_bytes)
    elif filename_lower.endswith('.docx'):
        return extract_text_from_docx(file_bytes)
    elif filename_lower.endswith('.doc'):
        # .doc files (old Word format) are not supported by python-docx
        # Return None and handle as fallback
        print(f"Warning: .doc files not supported, only .docx")
        return None
    else:
        print(f"Warning: Unsupported file type: {filename}")
        return None


def chunk_text(text: str, chunk_size: int = 1000, overlap: int = 200) -> list[str]:
    """
    Split text into chunks with optional overlap

    Args:
        text: Text to chunk
        chunk_size: Maximum characters per chunk
        overlap: Number of characters to overlap between chunks

    Returns:
        List of text chunks
    """
    if len(text) <= chunk_size:
        return [text]

    chunks = []
    start = 0

    while start < len(text):
        # Find the end of this chunk
        end = start + chunk_size

        # If this isn't the last chunk, try to break at a sentence or word boundary
        if end < len(text):
            # Look for sentence endings (. ! ?) within the last 100 chars
            last_period = text.rfind('.', start, end)
            last_exclaim = text.rfind('!', start, end)
            last_question = text.rfind('?', start, end)

            best_break = max(last_period, last_exclaim, last_question)

            # If no sentence break found, look for word boundary (space)
            if best_break == -1 or best_break < end - 100:
                last_space = text.rfind(' ', start, end)
                if last_space != -1:
                    best_break = last_space

            # Use the break point if found, otherwise use chunk_size
            if best_break != -1 and best_break > start:
                end = best_break + 1

        chunk = text[start:end].strip()
        if chunk:
            chunks.append(chunk)

        # Move start position, accounting for overlap
        start = end - overlap if end < len(text) else len(text)

    return chunks
