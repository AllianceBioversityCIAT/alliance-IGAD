"""
Proposals Router
"""

import json
import os
import uuid
from datetime import datetime
from typing import Any, Dict, List, Optional

import boto3
from fastapi import APIRouter, Depends, File, HTTPException, UploadFile
from fastapi.security import HTTPAuthorizationCredentials, HTTPBearer
from pydantic import BaseModel

from app.database.client import db_client
from app.middleware.auth_middleware import AuthMiddleware
from app.shared.ai.bedrock_service import BedrockService
from app.tools.admin.prompts_manager.service import PromptService
from app.tools.proposal_writer.proposal_template_generation.service import (
    ProposalTemplateGenerator,
)

router = APIRouter(prefix="/api/proposals", tags=["proposals"])
security = HTTPBearer()
auth_middleware = AuthMiddleware()

# Initialize Lambda client
lambda_client = boto3.client("lambda")


# Pydantic models
class ProposalCreate(BaseModel):
    title: str
    description: str = ""
    template_id: Optional[str] = None


class ProposalUpdate(BaseModel):
    title: Optional[str] = None
    description: Optional[str] = None
    status: Optional[str] = None
    uploaded_files: Optional[Dict[str, List[str]]] = None
    text_inputs: Optional[Dict[str, str]] = None
    metadata: Optional[Dict[str, Any]] = None


class AIGenerateRequest(BaseModel):
    section_id: str
    context_data: Optional[Dict[str, Any]] = None


class AIImproveRequest(BaseModel):
    section_id: str
    improvement_type: str = "general"


class PromptWithCategoriesRequest(BaseModel):
    prompt_id: str
    categories: List[str]


class ConceptEvaluationUpdate(BaseModel):
    selected_sections: List[Dict[str, Any]]
    user_comments: Optional[Dict[str, str]] = None


class TemplateGenerationRequest(BaseModel):
    selected_sections: Optional[List[str]] = None  # List of section titles to include
    user_comments: Optional[Dict[str, str]] = None  # Dict of {section_title: comment}


async def get_current_user(
    credentials: HTTPAuthorizationCredentials = Depends(security),
):
    """Get current user from token"""
    return auth_middleware.verify_token(credentials)


def generate_proposal_code() -> str:
    """Generate unique proposal code in format PROP-YYYYMMDD-XXXX"""
    now = datetime.utcnow()
    date_str = now.strftime("%Y%m%d")
    random_suffix = str(uuid.uuid4())[:4].upper()
    return f"PROP-{date_str}-{random_suffix}"


@router.post("")
async def create_proposal(proposal: ProposalCreate, user=Depends(get_current_user)):
    """Create a new proposal - only one draft allowed per user"""
    try:
        # Check if user already has a draft proposal
        existing_proposals = await db_client.query_items(
            pk=f"USER#{user.get('user_id')}",
            index_name="GSI1",
            scan_index_forward=False,
        )

        # Find existing draft
        existing_draft = None
        for prop in existing_proposals:
            if prop.get("status") == "draft":
                existing_draft = prop
                break

        # If draft exists, return it instead of creating new one
        if existing_draft:
            response_proposal = {
                k: v
                for k, v in existing_draft.items()
                if k not in ["PK", "SK", "GSI1PK", "GSI1SK"]
            }
            return {
                "proposal": response_proposal,
                "message": "Returning existing draft proposal",
            }

        # Create new proposal if no draft exists
        proposal_id = str(uuid.uuid4())
        proposal_code = generate_proposal_code()

        now = datetime.utcnow().isoformat()

        new_proposal = {
            "PK": f"PROPOSAL#{proposal_code}",
            "SK": "METADATA",
            "id": proposal_id,
            "proposalCode": proposal_code,
            "title": proposal.title,
            "description": proposal.description,
            "template_id": proposal.template_id,
            "status": "draft",
            "created_at": now,
            "updated_at": now,
            "user_id": user.get("user_id"),
            "user_email": user.get("email"),
            "user_name": user.get("name"),
            "uploaded_files": {},
            "text_inputs": {},
            "metadata": {},
            "GSI1PK": f"USER#{user.get('user_id')}",
            "GSI1SK": f"PROPOSAL#{now}",
        }

        await db_client.put_item(new_proposal)

        # Return proposal without DynamoDB keys
        response_proposal = {
            k: v
            for k, v in new_proposal.items()
            if k not in ["PK", "SK", "GSI1PK", "GSI1SK"]
        }

        return {"proposal": response_proposal}
    except Exception as e:
        raise HTTPException(
            status_code=500, detail=f"Failed to create proposal: {str(e)}"
        )


@router.get("")
async def get_proposals(user=Depends(get_current_user)):
    """Get all proposals for the current user"""
    try:
        # Query using GSI1 to get all proposals for this user
        items = await db_client.query_items(
            pk=f"USER#{user.get('user_id')}",
            index_name="GSI1",
            scan_index_forward=False,  # Most recent first
        )

        # Remove DynamoDB keys from response
        proposals = []
        for item in items:
            proposal = {
                k: v
                for k, v in item.items()
                if k not in ["PK", "SK", "GSI1PK", "GSI1SK"]
            }
            proposals.append(proposal)

        return {"proposals": proposals}
    except Exception as e:
        raise HTTPException(
            status_code=500, detail=f"Failed to fetch proposals: {str(e)}"
        )


@router.get("/{proposal_id}")
async def get_proposal(proposal_id: str, user=Depends(get_current_user)):
    """Get a specific proposal by ID or proposal code"""
    try:
        # Check if it's a proposal code (PROP-YYYYMMDD-XXXX) or UUID
        if proposal_id.startswith("PROP-"):
            pk = f"PROPOSAL#{proposal_id}"
        else:
            # Need to query by ID - use GSI to find it
            items = await db_client.query_items(
                pk=f"USER#{user.get('user_id')}", index_name="GSI1"
            )

            # Find the proposal with matching ID
            proposal_item = None
            for item in items:
                if item.get("id") == proposal_id:
                    proposal_item = item
                    break

            if not proposal_item:
                raise HTTPException(status_code=404, detail="Proposal not found")

            pk = proposal_item["PK"]

        # Get the proposal
        proposal = await db_client.get_item(pk=pk, sk="METADATA")

        if not proposal:
            raise HTTPException(status_code=404, detail="Proposal not found")

        # Verify ownership
        if proposal.get("user_id") != user.get("user_id"):
            raise HTTPException(status_code=403, detail="Access denied")

        # Remove DynamoDB keys from response
        response_proposal = {
            k: v
            for k, v in proposal.items()
            if k not in ["PK", "SK", "GSI1PK", "GSI1SK"]
        }

        return response_proposal
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(
            status_code=500, detail=f"Failed to fetch proposal: {str(e)}"
        )


@router.put("/{proposal_id}")
async def update_proposal(
    proposal_id: str, proposal_update: ProposalUpdate, user=Depends(get_current_user)
):
    """Update a proposal"""
    try:
        # First get the proposal to verify ownership and get PK
        if proposal_id.startswith("PROP-"):
            pk = f"PROPOSAL#{proposal_id}"
        else:
            # Query to find proposal by ID
            items = await db_client.query_items(
                pk=f"USER#{user.get('user_id')}", index_name="GSI1"
            )

            proposal_item = None
            for item in items:
                if item.get("id") == proposal_id:
                    proposal_item = item
                    break

            if not proposal_item:
                raise HTTPException(status_code=404, detail="Proposal not found")

            pk = proposal_item["PK"]

        # Get existing proposal
        existing_proposal = await db_client.get_item(pk=pk, sk="METADATA")

        if not existing_proposal:
            raise HTTPException(status_code=404, detail="Proposal not found")

        if existing_proposal.get("user_id") != user.get("user_id"):
            raise HTTPException(status_code=403, detail="Access denied")

        # Build update expression
        update_data = proposal_update.dict(exclude_unset=True)
        if not update_data:
            return {
                k: v
                for k, v in existing_proposal.items()
                if k not in ["PK", "SK", "GSI1PK", "GSI1SK"]
            }

        update_expression_parts = []
        expression_attribute_values = {}
        expression_attribute_names = {}

        for i, (field, value) in enumerate(update_data.items()):
            attr_name = f"#attr{i}"
            attr_value = f":val{i}"
            update_expression_parts.append(f"{attr_name} = {attr_value}")
            expression_attribute_names[attr_name] = field
            expression_attribute_values[attr_value] = value

        # Always update updated_at
        update_expression_parts.append("#updated_at = :updated_at")
        expression_attribute_names["#updated_at"] = "updated_at"
        expression_attribute_values[":updated_at"] = datetime.utcnow().isoformat()

        update_expression = "SET " + ", ".join(update_expression_parts)

        updated_proposal = await db_client.update_item(
            pk=pk,
            sk="METADATA",
            update_expression=update_expression,
            expression_attribute_values=expression_attribute_values,
            expression_attribute_names=expression_attribute_names,
        )

        # Remove DynamoDB keys from response
        response_proposal = {
            k: v
            for k, v in updated_proposal.items()
            if k not in ["PK", "SK", "GSI1PK", "GSI1SK"]
        }

        return response_proposal
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(
            status_code=500, detail=f"Failed to update proposal: {str(e)}"
        )


@router.delete("/{proposal_id}")
async def delete_proposal(proposal_id: str, user=Depends(get_current_user)):
    """Delete a proposal"""
    try:
        # First get the proposal to verify ownership and get PK
        if proposal_id.startswith("PROP-"):
            pk = f"PROPOSAL#{proposal_id}"
        else:
            # Query to find proposal by ID
            items = await db_client.query_items(
                pk=f"USER#{user.get('user_id')}", index_name="GSI1"
            )

            proposal_item = None
            for item in items:
                if item.get("id") == proposal_id:
                    proposal_item = item
                    break

            if not proposal_item:
                raise HTTPException(status_code=404, detail="Proposal not found")

            pk = proposal_item["PK"]

        # Get and verify ownership
        proposal = await db_client.get_item(pk=pk, sk="METADATA")

        if not proposal:
            raise HTTPException(status_code=404, detail="Proposal not found")

        if proposal.get("user_id") != user.get("user_id"):
            raise HTTPException(status_code=403, detail="Access denied")

        # Get proposal code for S3 cleanup
        proposal_code = proposal.get("proposalCode", proposal_id)

        print(f"🗑️  Starting cleanup for proposal: {proposal_code}")

        # ========== 1. DELETE S3 VECTORS ==========
        print("🔄 Deleting vectors from S3 Vectors...")
        try:
            from app.shared.vectors.service import VectorEmbeddingsService

            vector_service = VectorEmbeddingsService()

            vector_deleted = vector_service.delete_proposal_vectors(proposal_code)
            if vector_deleted:
                print(f"✅ Deleted vectors for {proposal_code}")
            else:
                print(f"⚠️  No vectors found or deletion failed for {proposal_code}")
        except Exception as vector_error:
            print(f"⚠️  Vector deletion error (non-critical): {str(vector_error)}")

        # ========== 2. DELETE S3 FILES ==========
        print("🔄 Deleting S3 documents...")
        try:
            import os

            from app.utils.aws_session import get_aws_session

            session = get_aws_session()
            s3_client = session.client("s3")
            bucket = os.environ.get("PROPOSALS_BUCKET")

            if bucket:
                # List all objects under the proposal folder
                prefix = f"{proposal_code}/"
                response = s3_client.list_objects_v2(Bucket=bucket, Prefix=prefix)

                if "Contents" in response:
                    objects_to_delete = [
                        {"Key": obj["Key"]} for obj in response["Contents"]
                    ]

                    if objects_to_delete:
                        s3_client.delete_objects(
                            Bucket=bucket, Delete={"Objects": objects_to_delete}
                        )
                        print(
                            f"✅ Deleted {len(objects_to_delete)} S3 objects for {proposal_code}"
                        )
                    else:
                        print(f"ℹ️  No S3 objects found for {proposal_code}")
                else:
                    print(f"ℹ️  No S3 objects found for {proposal_code}")
            else:
                print("⚠️  S3 bucket not configured")
        except Exception as s3_error:
            print(f"⚠️  S3 deletion error (non-critical): {str(s3_error)}")

        # ========== 3. DELETE DYNAMODB METADATA ==========
        print("🔄 Deleting DynamoDB metadata...")
        await db_client.delete_item(pk=pk, sk="METADATA")
        print(f"✅ Deleted DynamoDB metadata for {proposal_code}")

        print(f"✅ Proposal {proposal_code} deleted successfully with all resources")

        return {
            "message": "Proposal deleted successfully",
            "proposal_code": proposal_code,
            "cleanup_summary": {
                "vectors_deleted": "attempted",
                "s3_files_deleted": "attempted",
                "dynamodb_deleted": "completed",
            },
        }
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(
            status_code=500, detail=f"Failed to delete proposal: {str(e)}"
        )


@router.post("/{proposal_id}/generate")
async def generate_ai_content(
    proposal_id: str, request: AIGenerateRequest, user=Depends(get_current_user)
):
    """Generate AI content for a proposal section"""
    try:
        # First verify the proposal exists and user has access
        if proposal_id.startswith("PROP-"):
            pk = f"PROPOSAL#{proposal_id}"
        else:
            items = await db_client.query_items(
                pk=f"USER#{user.get('user_id')}", index_name="GSI1"
            )

            proposal_item = None
            for item in items:
                if item.get("id") == proposal_id:
                    proposal_item = item
                    break

            if not proposal_item:
                raise HTTPException(status_code=404, detail="Proposal not found")

            pk = proposal_item["PK"]

        proposal = await db_client.get_item(pk=pk, sk="METADATA")

        if not proposal:
            raise HTTPException(status_code=404, detail="Proposal not found")

        if proposal.get("user_id") != user.get("user_id"):
            raise HTTPException(status_code=403, detail="Access denied")

        # Initialize Bedrock service
        bedrock_service = BedrockService()

        # Generate content using AI
        generated_content = await bedrock_service.generate_content(
            section_id=request.section_id, context_data=request.context_data or {}
        )

        return {"generated_content": generated_content}
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"AI generation failed: {str(e)}")


@router.put("/{proposal_id}/concept-evaluation")
async def update_concept_evaluation(
    proposal_id: str, update: ConceptEvaluationUpdate, user=Depends(get_current_user)
):
    """Update concept evaluation with user's section selections and comments"""
    try:
        # Verify proposal exists and user has access
        if proposal_id.startswith("PROP-"):
            pk = f"PROPOSAL#{proposal_id}"
        else:
            items = await db_client.query_items(
                pk=f"USER#{user.get('user_id')}", index_name="GSI1"
            )

            proposal_item = None
            for item in items:
                if item.get("id") == proposal_id:
                    proposal_item = item
                    break

            if not proposal_item:
                raise HTTPException(status_code=404, detail="Proposal not found")

            pk = proposal_item["PK"]

        # Get existing proposal with concept_analysis
        proposal = await db_client.get_item(pk=pk, sk="METADATA")

        if not proposal:
            raise HTTPException(status_code=404, detail="Proposal not found")

        if proposal.get("user_id") != user.get("user_id"):
            raise HTTPException(status_code=403, detail="Access denied")

        # Get concept_analysis from proposal
        concept_analysis = proposal.get("concept_analysis")
        if not concept_analysis:
            raise HTTPException(
                status_code=404,
                detail="Concept analysis not found. Complete Step 1 first.",
            )

        print("=" * 80)
        print("🔍 UPDATE CONCEPT EVALUATION - Starting")
        print(f"📊 concept_analysis keys: {list(concept_analysis.keys())}")
        print(f"📝 Received {len(update.selected_sections)} sections from frontend")
        print("=" * 80)

        # Update the sections with user selections and comments
        # Handle both nested and non-nested concept_analysis structures
        if "concept_analysis" in concept_analysis:
            # Nested structure (from DynamoDB)
            inner_analysis = concept_analysis["concept_analysis"]
        else:
            inner_analysis = concept_analysis

        # Get sections_needing_elaboration (the correct field name)
        sections = inner_analysis.get("sections_needing_elaboration", [])

        if sections:
            # Create a map of section titles to user selections
            user_selections = {s["title"]: s for s in update.selected_sections}

            # Update each section with user's selection and comments
            for section in sections:
                # Section title can be in "section" or "title" field
                title = section.get("section", section.get("title", ""))
                if title in user_selections:
                    user_section = user_selections[title]
                    section["selected"] = user_section.get("selected", True)

                    # Add user comments if provided
                    if update.user_comments and title in update.user_comments:
                        section["user_comment"] = update.user_comments[title]
                else:
                    # If section not in user_selections, mark as NOT selected
                    section["selected"] = False

            print(f"✅ Updated {len(sections)} sections with user selections")
            for section in sections:
                title = section.get("section", section.get("title", "Unknown"))
                selected = section.get("selected", False)
                print(f"   • {title}: selected={selected}")
            print("=" * 80)

        # Update concept_analysis back in the proposal METADATA record
        await db_client.update_item(
            pk=pk,
            sk="METADATA",
            update_expression="SET concept_analysis = :analysis, updated_at = :updated",
            expression_attribute_values={
                ":analysis": concept_analysis,
                ":updated": datetime.utcnow().isoformat(),
            },
        )

        return {
            "status": "success",
            "message": "Concept evaluation updated successfully",
            "concept_evaluation": concept_analysis,
        }
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(
            status_code=500, detail=f"Failed to update concept evaluation: {str(e)}"
        )


@router.post("/{proposal_id}/improve")
async def improve_ai_content(
    proposal_id: str, request: AIImproveRequest, user=Depends(get_current_user)
):
    """Improve existing content using AI"""
    try:
        # First verify the proposal exists and user has access
        if proposal_id.startswith("PROP-"):
            pk = f"PROPOSAL#{proposal_id}"
        else:
            items = await db_client.query_items(
                pk=f"USER#{user.get('user_id')}", index_name="GSI1"
            )

            proposal_item = None
            for item in items:
                if item.get("id") == proposal_id:
                    proposal_item = item
                    break

            if not proposal_item:
                raise HTTPException(status_code=404, detail="Proposal not found")

            pk = proposal_item["PK"]

        proposal = await db_client.get_item(pk=pk, sk="METADATA")

        if not proposal:
            raise HTTPException(status_code=404, detail="Proposal not found")

        if proposal.get("user_id") != user.get("user_id"):
            raise HTTPException(status_code=403, detail="Access denied")

        # Initialize Bedrock service
        bedrock_service = BedrockService()

        # Improve content using AI
        improved_content = await bedrock_service.improve_content(
            section_id=request.section_id, improvement_type=request.improvement_type
        )

        return {"improved_content": improved_content}
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"AI improvement failed: {str(e)}")


@router.post("/{proposal_id}/analyze-rfp")
async def analyze_rfp(proposal_id: str, user=Depends(get_current_user)):
    """
    Start RFP analysis (async - returns immediately with status)
    Frontend should poll GET /{proposal_id}/analysis-status for completion
    """
    try:
        # Verify proposal ownership
        if proposal_id.startswith("PROP-"):
            pk = f"PROPOSAL#{proposal_id}"
            proposal_code = proposal_id
        else:
            items = await db_client.query_items(
                pk=f"USER#{user.get('user_id')}", index_name="GSI1"
            )

            proposal_item = None
            for item in items:
                if item.get("id") == proposal_id:
                    proposal_item = item
                    break

            if not proposal_item:
                raise HTTPException(status_code=404, detail="Proposal not found")

            pk = proposal_item["PK"]
            proposal_code = proposal_item.get("proposalCode", proposal_id)

        proposal = await db_client.get_item(pk=pk, sk="METADATA")

        if not proposal:
            raise HTTPException(status_code=404, detail="Proposal not found")

        if proposal.get("user_id") != user.get("user_id"):
            raise HTTPException(status_code=403, detail="Access denied")

        # Check if RFP analysis already exists (no re-analysis)
        if proposal.get("rfp_analysis"):
            return {
                "status": "completed",
                "rfp_analysis": proposal.get("rfp_analysis"),
                "message": "RFP already analyzed",
                "cached": True,
            }

        # Check if analysis is already in progress
        analysis_status = proposal.get("analysis_status_rfp")
        if analysis_status == "processing":
            return {
                "status": "processing",
                "message": "Analysis already in progress",
                "started_at": proposal.get("rfp_analysis_started_at"),
            }

        # Update status to processing
        await db_client.update_item(
            pk=pk,
            sk="METADATA",
            update_expression="SET analysis_status_rfp = :status, rfp_analysis_started_at = :started",
            expression_attribute_values={
                ":status": "processing",
                ":started": datetime.utcnow().isoformat(),
            },
        )

        # Get proposal code for Worker
        proposal_code = proposal.get("proposalCode")
        if not proposal_code:
            raise HTTPException(status_code=400, detail="Proposal code not found")

        # Invoke Worker Lambda asynchronously
        print(f"🚀 Invoking AnalysisWorkerFunction for proposal {proposal_code}")

        # Get worker function ARN from environment variable
        worker_function_arn = os.environ.get("WORKER_FUNCTION_NAME")
        if not worker_function_arn:
            raise Exception("WORKER_FUNCTION_NAME environment variable not set")

        print(f"📝 Worker function ARN: {worker_function_arn}")

        # Invoke async
        lambda_client.invoke(
            FunctionName=worker_function_arn,
            InvocationType="Event",  # Async invocation
            Payload=json.dumps(
                {
                    "proposal_id": proposal_code,  # Send proposal_code, not UUID
                    "analysis_type": "rfp",
                }
            ),
        )

        print(f"✅ Worker Lambda invoked successfully")

        return {
            "status": "processing",
            "message": "RFP analysis started. Poll /analysis-status for completion.",
            "started_at": datetime.utcnow().isoformat(),
        }

    except HTTPException:
        raise
    except Exception as e:
        import traceback

        error_details = traceback.format_exc()
        print(f"❌ ERROR in analyze_rfp endpoint:")
        print(error_details)
        raise HTTPException(status_code=500, detail=f"RFP analysis failed: {str(e)}")


@router.get("/{proposal_id}/analysis-status")
async def get_analysis_status(proposal_id: str, user=Depends(get_current_user)):
    """Poll for RFP analysis completion status"""
    try:
        # Verify proposal ownership
        if proposal_id.startswith("PROP-"):
            pk = f"PROPOSAL#{proposal_id}"
        else:
            items = await db_client.query_items(
                pk=f"USER#{user.get('user_id')}", index_name="GSI1"
            )

            proposal_item = None
            for item in items:
                if item.get("id") == proposal_id:
                    proposal_item = item
                    break

            if not proposal_item:
                raise HTTPException(status_code=404, detail="Proposal not found")

            pk = proposal_item["PK"]

        proposal = await db_client.get_item(pk=pk, sk="METADATA")

        if not proposal:
            raise HTTPException(status_code=404, detail="Proposal not found")

        if proposal.get("user_id") != user.get("user_id"):
            raise HTTPException(status_code=403, detail="Access denied")

        status = proposal.get("analysis_status_rfp", "not_started")

        if status == "completed":
            return {
                "status": "completed",
                "rfp_analysis": proposal.get("rfp_analysis"),
                "completed_at": proposal.get("rfp_analysis_completed_at"),
            }
        elif status == "failed":
            return {
                "status": "failed",
                "error": proposal.get("rfp_analysis_error", "Unknown error"),
            }
        elif status == "processing":
            return {
                "status": "processing",
                "started_at": proposal.get("rfp_analysis_started_at"),
            }
        else:
            return {"status": "not_started"}

    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to check status: {str(e)}")


class ConceptAnalysisRequest(BaseModel):
    force: bool = False


@router.post("/{proposal_id}/analyze-concept")
async def analyze_concept(
    proposal_id: str,
    request: Optional[ConceptAnalysisRequest] = None,
    user=Depends(get_current_user),
):
    """
    Start Concept analysis (async - returns immediately with status)
    Requires RFP analysis to be completed first.

    Args:
        force: If True, forces a new analysis even if one already exists.
              Use this when the concept document has been re-uploaded.
    """
    try:
        # Get proposal
        if proposal_id.startswith("PROP-"):
            pk = f"PROPOSAL#{proposal_id}"
            proposal_code = proposal_id
        else:
            items = await db_client.query_items(
                pk=f"USER#{user.get('user_id')}", index_name="GSI1"
            )

            proposal_item = None
            for item in items:
                if item.get("id") == proposal_id:
                    proposal_item = item
                    break

            if not proposal_item:
                raise HTTPException(status_code=404, detail="Proposal not found")

            pk = proposal_item["PK"]
            proposal_code = proposal_item.get("proposalCode", proposal_id)

        proposal = await db_client.get_item(pk=pk, sk="METADATA")

        if not proposal:
            raise HTTPException(status_code=404, detail="Proposal not found")

        if proposal.get("user_id") != user.get("user_id"):
            raise HTTPException(status_code=403, detail="Access denied")

        # Check if RFP analysis exists
        if not proposal.get("rfp_analysis"):
            raise HTTPException(
                status_code=400, detail="RFP analysis must be completed first"
            )

        # Get force flag from request body
        force = request.force if request else False

        # Check if concept analysis already exists (unless force=True)
        if proposal.get("concept_analysis") and not force:
            return {
                "status": "completed",
                "concept_analysis": proposal.get("concept_analysis"),
                "message": "Concept already analyzed",
                "cached": True,
            }

        # If force=True, clear existing analysis data
        if force and proposal.get("concept_analysis"):
            print(
                f"🔄 Force re-analysis requested, clearing existing concept analysis for {proposal_code}"
            )
            await db_client.update_item(
                pk=pk,
                sk="METADATA",
                update_expression="""
                    REMOVE concept_analysis,
                           concept_analysis_completed_at,
                           concept_analysis_error,
                           concept_evaluation,
                           concept_document_v2,
                           structure_workplan_analysis,
                           structure_workplan_completed_at,
                           structure_workplan_error
                    SET analysis_status_concept = :not_started,
                        updated_at = :updated
                """,
                expression_attribute_values={
                    ":not_started": "not_started",
                    ":updated": datetime.utcnow().isoformat(),
                },
            )

        # Check if analysis is already in progress (unless force=True)
        analysis_status = proposal.get("analysis_status_concept")
        if analysis_status == "processing" and not force:
            return {
                "status": "processing",
                "message": "Concept analysis already in progress",
                "started_at": proposal.get("concept_analysis_started_at"),
            }

        # Update status to processing
        await db_client.update_item(
            pk=pk,
            sk="METADATA",
            update_expression="SET analysis_status_concept = :status, concept_analysis_started_at = :started",
            expression_attribute_values={
                ":status": "processing",
                ":started": datetime.utcnow().isoformat(),
            },
        )

        # Invoke Worker Lambda asynchronously
        print(
            f"🚀 Invoking AnalysisWorkerFunction for concept analysis: {proposal_code}"
        )

        worker_function_arn = os.environ.get("WORKER_FUNCTION_NAME")
        if not worker_function_arn:
            raise Exception("WORKER_FUNCTION_NAME environment variable not set")

        lambda_client.invoke(
            FunctionName=worker_function_arn,
            InvocationType="Event",
            Payload=json.dumps(
                {"proposal_id": proposal_code, "analysis_type": "concept"}
            ),
        )

        print(f"✅ Concept analysis worker invoked successfully")

        return {
            "status": "processing",
            "message": "Concept analysis started. Poll /concept-status for completion.",
            "started_at": datetime.utcnow().isoformat(),
        }

    except HTTPException:
        raise
    except Exception as e:
        import traceback

        error_details = traceback.format_exc()
        print(f"❌ ERROR in analyze_concept endpoint:")
        print(error_details)
        raise HTTPException(
            status_code=500, detail=f"Concept analysis failed: {str(e)}"
        )


@router.get("/{proposal_id}/concept-status")
async def get_concept_status(proposal_id: str, user=Depends(get_current_user)):
    """Poll for Concept analysis completion status"""
    try:
        # Get proposal
        if proposal_id.startswith("PROP-"):
            pk = f"PROPOSAL#{proposal_id}"
        else:
            items = await db_client.query_items(
                pk=f"USER#{user.get('user_id')}", index_name="GSI1"
            )

            proposal_item = None
            for item in items:
                if item.get("id") == proposal_id:
                    proposal_item = item
                    break

            if not proposal_item:
                raise HTTPException(status_code=404, detail="Proposal not found")

            pk = proposal_item["PK"]

        proposal = await db_client.get_item(pk=pk, sk="METADATA")

        if not proposal:
            raise HTTPException(status_code=404, detail="Proposal not found")

        if proposal.get("user_id") != user.get("user_id"):
            raise HTTPException(status_code=403, detail="Access denied")

        status = proposal.get("analysis_status_concept", "not_started")

        if status == "completed":
            return {
                "status": "completed",
                "concept_analysis": proposal.get("concept_analysis"),
                "completed_at": proposal.get("concept_analysis_completed_at"),
            }
        elif status == "failed":
            return {
                "status": "failed",
                "error": proposal.get("concept_analysis_error", "Unknown error"),
            }
        elif status == "processing":
            return {
                "status": "processing",
                "started_at": proposal.get("concept_analysis_started_at"),
            }
        else:
            return {"status": "not_started"}

    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(
            status_code=500, detail=f"Failed to check concept status: {str(e)}"
        )


@router.post("/{proposal_id}/generate-concept-document")
async def generate_concept_document(
    proposal_id: str, concept_evaluation: dict, user=Depends(get_current_user)
):
    """
    Generate updated concept document based on RFP analysis and user's concept evaluation

    This endpoint:
    1. Validates proposal exists
    2. Sets status to "processing"
    3. Invokes Worker Lambda asynchronously
    4. Returns immediately with status

    Frontend should poll /concept-document-status for completion
    """
    try:
        user_id = user.get("user_id")
        print(f"🟢 Generate concept document request for proposal: {proposal_id}")

        # Get proposal
        all_proposals = await db_client.query_items(
            pk=f"USER#{user_id}", index_name="GSI1"
        )

        proposal = None
        for p in all_proposals:
            if p.get("id") == proposal_id:
                proposal = p
                break

        if not proposal:
            raise HTTPException(status_code=404, detail="Proposal not found")

        # Validate required data exists
        rfp_analysis = proposal.get("rfp_analysis")
        if not rfp_analysis:
            raise HTTPException(
                status_code=400, detail="RFP analysis not found. Complete Step 1 first."
            )

        concept_analysis = proposal.get("concept_analysis")
        if not concept_analysis:
            raise HTTPException(
                status_code=400,
                detail="Concept analysis not found. Complete Step 1 first.",
            )

        # 🔥 CRITICAL FIX: Build concept_evaluation from DynamoDB data (which has selected sections)
        # The concept_evaluation param from frontend is just a trigger, not the source of truth
        print("=" * 80)
        print("🔍 Building concept_evaluation from DynamoDB...")
        print(
            f"📊 concept_analysis keys: {list(concept_analysis.keys()) if isinstance(concept_analysis, dict) else 'NOT A DICT'}"
        )

        # Use the concept_analysis from DynamoDB (which has user selections)
        final_concept_evaluation = {
            "concept_analysis": concept_analysis,
            "status": "completed",
        }

        print(
            f"✅ Final concept_evaluation has {len(concept_analysis.get('sections_needing_elaboration', []))} sections"
        )
        print("=" * 80)

        # Update status to processing
        await db_client.update_item(
            pk=proposal["PK"],
            sk=proposal["SK"],
            update_expression="SET concept_document_status = :status, concept_document_started_at = :started",
            expression_attribute_values={
                ":status": "processing",
                ":started": datetime.utcnow().isoformat(),
            },
        )

        # Invoke Worker Lambda asynchronously
        worker_function = os.getenv("WORKER_FUNCTION_NAME")

        # Get proposal code for Worker (same as Step 1)
        proposal_code = proposal.get("proposalCode")
        if not proposal_code:
            raise HTTPException(status_code=400, detail="Proposal code not found")

        payload = {
            "analysis_type": "concept_document",
            "proposal_id": proposal_code,  # Send proposal_code, not UUID
            "user_id": user_id,
            "concept_evaluation": final_concept_evaluation,  # 🔥 Use DynamoDB data, not request param
        }

        print(
            f"📡 Invoking worker lambda for concept document generation: {proposal_id}"
        )

        lambda_client.invoke(
            FunctionName=worker_function,
            InvocationType="Event",  # Asynchronous
            Payload=json.dumps(payload),
        )

        print(f"✅ Worker lambda invoked successfully for {proposal_id}")

        # Return immediately
        return {
            "status": "processing",
            "message": "Concept document generation started. Poll /concept-document-status for updates.",
        }

    except HTTPException:
        raise
    except Exception as e:
        print(f"❌ Error starting concept document generation: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/{proposal_id}/concept-document-status")
async def get_concept_document_status(proposal_id: str, user=Depends(get_current_user)):
    """Get concept document generation status"""
    try:
        all_proposals = await db_client.query_items(
            pk=f"USER#{user.get('user_id')}", index_name="GSI1"
        )

        proposal = None
        for p in all_proposals:
            if p.get("id") == proposal_id:
                proposal = p
                break

        if not proposal:
            raise HTTPException(status_code=404, detail="Proposal not found")

        status = proposal.get("concept_document_status", "not_started")

        response = {
            "status": status,
            "started_at": proposal.get("concept_document_started_at"),
            "completed_at": proposal.get("concept_document_completed_at"),
        }

        if status == "completed":
            response["concept_document"] = proposal.get("concept_document_v2")
        elif status == "failed":
            response["error"] = proposal.get("concept_document_error")

        return response

    except Exception as e:
        print(f"❌ Error getting concept document status: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/{proposal_id}/concept-evaluation")
async def get_concept_evaluation(proposal_id: str, user=Depends(get_current_user)):
    """
    Get the saved concept_evaluation from DynamoDB
    """
    try:
        user_id = user.get("user_id")
        print(f"📥 Getting concept evaluation for proposal: {proposal_id}")

        # Get proposal
        all_proposals = await db_client.query_items(
            pk=f"USER#{user_id}", index_name="GSI1"
        )

        proposal = None
        for p in all_proposals:
            if p.get("id") == proposal_id:
                proposal = p
                break

        if not proposal:
            raise HTTPException(status_code=404, detail="Proposal not found")

        concept_evaluation = proposal.get("concept_evaluation")

        if not concept_evaluation:
            # Fallback to concept_analysis if concept_evaluation doesn't exist yet
            concept_evaluation = proposal.get("concept_analysis")

        if not concept_evaluation:
            raise HTTPException(status_code=404, detail="No concept evaluation found")

        print(f"✅ Concept evaluation retrieved for {proposal_id}")

        return {"concept_evaluation": concept_evaluation}

    except HTTPException:
        raise
    except Exception as e:
        print(f"❌ Error getting concept evaluation: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/{proposal_id}/analyze-step-1")
async def analyze_step_1(proposal_id: str, user=Depends(get_current_user)):
    """
    Step 1: RFP Analysis ONLY

    This endpoint:
    1. Triggers RFP analysis
    2. Returns immediately with status

    Frontend should poll /step-1-status for completion, then call /analyze-step-2
    """
    try:
        # Verify proposal ownership
        if proposal_id.startswith("PROP-"):
            pk = f"PROPOSAL#{proposal_id}"
            proposal_code = proposal_id
        else:
            items = await db_client.query_items(
                pk=f"USER#{user.get('user_id')}", index_name="GSI1"
            )

            proposal_item = None
            for item in items:
                if item.get("id") == proposal_id:
                    proposal_item = item
                    break

            if not proposal_item:
                raise HTTPException(status_code=404, detail="Proposal not found")

            pk = proposal_item["PK"]
            proposal_code = proposal_item.get("proposalCode", proposal_id)

        proposal = await db_client.get_item(pk=pk, sk="METADATA")

        if not proposal:
            raise HTTPException(status_code=404, detail="Proposal not found")

        if proposal.get("user_id") != user.get("user_id"):
            raise HTTPException(status_code=403, detail="Access denied")

        # Get proposal code for Worker
        proposal_code = proposal.get("proposalCode")
        if not proposal_code:
            raise HTTPException(status_code=400, detail="Proposal code not found")

        # Get worker function ARN from environment variable
        worker_function_arn = os.environ.get("WORKER_FUNCTION_NAME")
        if not worker_function_arn:
            raise Exception("WORKER_FUNCTION_NAME environment variable not set")

        # Track which analyses to start
        analyses_started = []

        # ========== RFP ANALYSIS ==========
        # Check if RFP analysis already exists (no re-analysis)
        if proposal.get("rfp_analysis"):
            print(f"✓ RFP analysis already completed for {proposal_code}")
            analyses_started.append(
                {"type": "rfp", "status": "completed", "cached": True}
            )
        else:
            # Check if RFP analysis is already in progress
            rfp_status = proposal.get("analysis_status_rfp")
            if rfp_status == "processing":
                print(f"⏳ RFP analysis already in progress for {proposal_code}")
                analyses_started.append(
                    {"type": "rfp", "status": "processing", "already_running": True}
                )
            else:
                # Start RFP analysis
                await db_client.update_item(
                    pk=pk,
                    sk="METADATA",
                    update_expression="SET analysis_status_rfp = :status, rfp_analysis_started_at = :started",
                    expression_attribute_values={
                        ":status": "processing",
                        ":started": datetime.utcnow().isoformat(),
                    },
                )

                print(f"🚀 Invoking RFP analysis for {proposal_code}")
                lambda_client.invoke(
                    FunctionName=worker_function_arn,
                    InvocationType="Event",  # Async invocation
                    Payload=json.dumps(
                        {"proposal_id": proposal_code, "analysis_type": "rfp"}
                    ),
                )

                analyses_started.append(
                    {"type": "rfp", "status": "processing", "started": True}
                )

        print(f"✅ Step 1 (RFP) analysis triggered successfully for {proposal_code}")

        return {
            "status": "processing",
            "message": "Step 1 (RFP) analysis started. Poll /step-1-status for completion, then call /analyze-step-2.",
            "analyses": analyses_started,
            "started_at": datetime.utcnow().isoformat(),
        }

    except HTTPException:
        raise
    except Exception as e:
        import traceback

        error_details = traceback.format_exc()
        print(f"❌ ERROR in analyze_step_1 endpoint:")
        print(error_details)
        raise HTTPException(status_code=500, detail=f"Step 1 analysis failed: {str(e)}")


@router.post("/{proposal_id}/analyze-step-2")
async def analyze_step_2(proposal_id: str, user=Depends(get_current_user)):
    """
    Step 2: Reference Proposals + Existing Work Analysis

    Prerequisites: Step 1 (RFP) must be completed with semantic_query

    This endpoint:
    1. Verifies RFP analysis is completed
    2. Triggers Reference Proposals analysis (semantic search)
    3. Triggers Existing Work analysis (semantic search) - TODO
    4. Returns immediately with status

    Frontend should poll for completion before calling /analyze-step-3
    """
    try:
        # Verify proposal ownership
        if proposal_id.startswith("PROP-"):
            pk = f"PROPOSAL#{proposal_id}"
            proposal_code = proposal_id
        else:
            items = await db_client.query_items(
                pk=f"USER#{user.get('user_id')}", index_name="GSI1"
            )

            proposal_item = None
            for item in items:
                if item.get("id") == proposal_id:
                    proposal_item = item
                    break

            if not proposal_item:
                raise HTTPException(status_code=404, detail="Proposal not found")

            pk = proposal_item["PK"]
            proposal_code = proposal_item.get("proposalCode", proposal_id)

        proposal = await db_client.get_item(pk=pk, sk="METADATA")

        if not proposal:
            raise HTTPException(status_code=404, detail="Proposal not found")

        if proposal.get("user_id") != user.get("user_id"):
            raise HTTPException(status_code=403, detail="Access denied")

        proposal_code = proposal.get("proposalCode")
        if not proposal_code:
            raise HTTPException(status_code=400, detail="Proposal code not found")

        # ========== PREREQUISITE CHECK: RFP MUST BE COMPLETED ==========
        rfp_analysis = proposal.get("rfp_analysis")
        if not rfp_analysis:
            raise HTTPException(
                status_code=400,
                detail="Step 1 (RFP analysis) must be completed before Step 2.",
            )

        # Handle nested structure (rfp_analysis.rfp_analysis.semantic_query)
        # or flat structure (rfp_analysis.semantic_query)
        semantic_query = rfp_analysis.get("semantic_query")
        if not semantic_query and "rfp_analysis" in rfp_analysis:
            # Try nested structure
            nested_rfp = rfp_analysis.get("rfp_analysis", {})
            semantic_query = nested_rfp.get("semantic_query")
            if semantic_query:
                print(
                    f"ℹ️  Found semantic_query in nested structure for {proposal_code}"
                )

        if not semantic_query:
            # Log the structure for debugging
            print(f"❌ No semantic_query found in RFP analysis for {proposal_code}")
            print(f"   Available keys in rfp_analysis: {list(rfp_analysis.keys())}")
            if "rfp_analysis" in rfp_analysis:
                nested_keys = list(rfp_analysis.get("rfp_analysis", {}).keys())
                print(f"   Available keys in nested rfp_analysis: {nested_keys}")
            raise HTTPException(
                status_code=400,
                detail="Step 1 (RFP analysis) must be completed with semantic_query before Step 2.",
            )

        print(f"✓ RFP analysis completed with semantic_query for {proposal_code}")
        print(f"  Semantic query: {semantic_query[:100]}...")

        # Get worker function name
        worker_function_name = os.environ.get("WORKER_FUNCTION_NAME")
        if not worker_function_name:
            raise Exception("WORKER_FUNCTION_NAME environment variable not set")

        analyses_started = []

        # ========== REFERENCE PROPOSALS ANALYSIS ==========
        if proposal.get("reference_proposal_analysis"):
            print(f"✓ Reference Proposals already completed for {proposal_code}")
            analyses_started.append(
                {"type": "reference_proposals", "status": "completed", "cached": True}
            )
        else:
            ref_status = proposal.get("analysis_status_reference_proposals")
            if ref_status == "processing":
                print(f"⏳ Reference Proposals already in progress for {proposal_code}")
                analyses_started.append(
                    {
                        "type": "reference_proposals",
                        "status": "processing",
                        "already_running": True,
                    }
                )
            else:
                print(f"🚀 Invoking Reference Proposals analysis for {proposal_code}")

                await db_client.update_item(
                    pk=pk,
                    sk="METADATA",
                    update_expression="SET analysis_status_reference_proposals = :status, reference_proposals_started_at = :started",
                    expression_attribute_values={
                        ":status": "processing",
                        ":started": datetime.utcnow().isoformat(),
                    },
                )

                lambda_client.invoke(
                    FunctionName=worker_function_name,
                    InvocationType="Event",
                    Payload=json.dumps(
                        {
                            "proposal_id": proposal_code,
                            "analysis_type": "reference_proposals",
                        }
                    ),
                )

                analyses_started.append(
                    {
                        "type": "reference_proposals",
                        "status": "processing",
                        "started": True,
                    }
                )

        # ========== EXISTING WORK ANALYSIS ==========
        if proposal.get("existing_work_analysis"):
            print(f"✓ Existing Work already completed for {proposal_code}")
            analyses_started.append(
                {"type": "existing_work", "status": "completed", "cached": True}
            )
        else:
            existing_work_status = proposal.get("analysis_status_existing_work")
            if existing_work_status == "processing":
                print(f"⏳ Existing Work already in progress for {proposal_code}")
                analyses_started.append(
                    {
                        "type": "existing_work",
                        "status": "processing",
                        "already_running": True,
                    }
                )
            else:
                print(f"🚀 Invoking Existing Work analysis for {proposal_code}")

                await db_client.update_item(
                    pk=pk,
                    sk="METADATA",
                    update_expression="SET analysis_status_existing_work = :status, existing_work_started_at = :started",
                    expression_attribute_values={
                        ":status": "processing",
                        ":started": datetime.utcnow().isoformat(),
                    },
                )

                lambda_client.invoke(
                    FunctionName=worker_function_name,
                    InvocationType="Event",
                    Payload=json.dumps(
                        {"proposal_id": proposal_code, "analysis_type": "existing_work"}
                    ),
                )

                analyses_started.append(
                    {"type": "existing_work", "status": "processing", "started": True}
                )

        print(f"✅ Step 2 analysis triggered successfully for {proposal_code}")

        return {
            "status": "processing",
            "message": "Step 2 (Reference Proposals + Existing Work) analysis started. Poll for completion before calling /analyze-step-3.",
            "analyses": analyses_started,
            "started_at": datetime.utcnow().isoformat(),
        }

    except HTTPException:
        raise
    except Exception as e:
        import traceback

        error_details = traceback.format_exc()
        print(f"❌ ERROR in analyze_step_2 endpoint:")
        print(error_details)
        raise HTTPException(status_code=500, detail=f"Step 2 analysis failed: {str(e)}")


@router.get("/{proposal_id}/step-1-status")
async def get_step_1_status(proposal_id: str, user=Depends(get_current_user)):
    """
    Poll for Step 1 analysis completion status
    Returns ONLY RFP analysis status (Reference Proposals moved to Step 2)
    """
    try:
        # Verify proposal ownership
        if proposal_id.startswith("PROP-"):
            pk = f"PROPOSAL#{proposal_id}"
        else:
            items = await db_client.query_items(
                pk=f"USER#{user.get('user_id')}", index_name="GSI1"
            )

            proposal_item = None
            for item in items:
                if item.get("id") == proposal_id:
                    proposal_item = item
                    break

            if not proposal_item:
                raise HTTPException(status_code=404, detail="Proposal not found")

            pk = proposal_item["PK"]

        proposal = await db_client.get_item(pk=pk, sk="METADATA")

        if not proposal:
            raise HTTPException(status_code=404, detail="Proposal not found")

        if proposal.get("user_id") != user.get("user_id"):
            raise HTTPException(status_code=403, detail="Access denied")

        # Get status of RFP analysis ONLY
        rfp_status = proposal.get("analysis_status_rfp", "not_started")

        # Build response
        response = {"rfp_analysis": {"status": rfp_status}}

        # Add RFP details
        if rfp_status == "completed":
            response["rfp_analysis"]["data"] = proposal.get("rfp_analysis")
            response["rfp_analysis"]["completed_at"] = proposal.get(
                "rfp_analysis_completed_at"
            )
        elif rfp_status == "failed":
            response["rfp_analysis"]["error"] = proposal.get(
                "rfp_analysis_error", "Unknown error"
            )
        elif rfp_status == "processing":
            response["rfp_analysis"]["started_at"] = proposal.get(
                "rfp_analysis_started_at"
            )

        # Overall status is just RFP status (Step 1 = RFP only)
        response["overall_status"] = rfp_status

        return response

    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(
            status_code=500, detail=f"Failed to check Step 1 status: {str(e)}"
        )


@router.get("/{proposal_id}/step-2-status")
async def get_step_2_status(proposal_id: str, user=Depends(get_current_user)):
    """
    Poll for Step 2 analysis completion status
    Returns combined status of Reference Proposals and Existing Work analyses

    Overall status is 'completed' only when BOTH analyses are done.
    Both analyses run in parallel after Step 1 (RFP) completes.
    """
    try:
        # Verify proposal ownership
        if proposal_id.startswith("PROP-"):
            pk = f"PROPOSAL#{proposal_id}"
        else:
            items = await db_client.query_items(
                pk=f"USER#{user.get('user_id')}", index_name="GSI1"
            )

            proposal_item = None
            for item in items:
                if item.get("id") == proposal_id:
                    proposal_item = item
                    break

            if not proposal_item:
                raise HTTPException(status_code=404, detail="Proposal not found")

            pk = proposal_item["PK"]

        proposal = await db_client.get_item(pk=pk, sk="METADATA")

        if not proposal:
            raise HTTPException(status_code=404, detail="Proposal not found")

        if proposal.get("user_id") != user.get("user_id"):
            raise HTTPException(status_code=403, detail="Access denied")

        # Get status of Step 2 analyses
        ref_status = proposal.get("analysis_status_reference_proposals", "not_started")
        existing_work_status = proposal.get(
            "analysis_status_existing_work", "not_started"
        )

        # Build response
        response = {
            "reference_proposals_analysis": {"status": ref_status},
            "existing_work_analysis": {"status": existing_work_status},
        }

        # Add Reference Proposals details
        if ref_status == "completed":
            response["reference_proposals_analysis"]["data"] = proposal.get(
                "reference_proposal_analysis"
            )
            response["reference_proposals_analysis"]["completed_at"] = proposal.get(
                "reference_proposals_completed_at"
            )
        elif ref_status == "failed":
            response["reference_proposals_analysis"]["error"] = proposal.get(
                "reference_proposals_error", "Unknown error"
            )
        elif ref_status == "processing":
            response["reference_proposals_analysis"]["started_at"] = proposal.get(
                "reference_proposals_started_at"
            )

        # Add Existing Work details
        if existing_work_status == "completed":
            response["existing_work_analysis"]["data"] = proposal.get(
                "existing_work_analysis"
            )
            response["existing_work_analysis"]["completed_at"] = proposal.get(
                "existing_work_completed_at"
            )
        elif existing_work_status == "failed":
            response["existing_work_analysis"]["error"] = proposal.get(
                "existing_work_error", "Unknown error"
            )
        elif existing_work_status == "processing":
            response["existing_work_analysis"]["started_at"] = proposal.get(
                "existing_work_started_at"
            )

        # Determine overall status
        # Both Reference Proposals AND Existing Work must complete for Step 2 to be done
        if ref_status == "completed" and existing_work_status == "completed":
            response["overall_status"] = "completed"
        elif ref_status == "failed" or existing_work_status == "failed":
            response["overall_status"] = "failed"
        elif ref_status == "processing" or existing_work_status == "processing":
            response["overall_status"] = "processing"
        else:
            response["overall_status"] = "not_started"

        return response

    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(
            status_code=500, detail=f"Failed to check Step 2 status: {str(e)}"
        )


# ==================== HELPER FUNCTIONS ====================


async def _verify_proposal_access(
    proposal_id: str, user: Dict[str, Any]
) -> tuple[str, str, Dict[str, Any]]:
    """
    Verify user has access to proposal and return metadata.

    This helper function reduces code duplication across endpoints by centralizing
    the logic for:
    1. Resolving proposal_id (UUID vs PROP-CODE)
    2. Verifying proposal ownership
    3. Loading proposal metadata

    Args:
        proposal_id: Proposal UUID or code (PROP-YYYYMMDD-XXXX)
        user: Authenticated user dict from get_current_user()

    Returns:
        Tuple of (pk, proposal_code, proposal_data):
            - pk: DynamoDB partition key (PROPOSAL#{code})
            - proposal_code: Resolved proposal code (PROP-YYYYMMDD-XXXX)
            - proposal_data: Full proposal metadata dict

    Raises:
        HTTPException(404): If proposal not found
        HTTPException(403): If user doesn't own the proposal

    Example:
        pk, proposal_code, proposal = await _verify_proposal_access(proposal_id, user)
        # Now you can use pk, proposal_code, and proposal
    """
    # Resolve proposal ID (UUID vs code)
    if proposal_id.startswith("PROP-"):
        pk = f"PROPOSAL#{proposal_id}"
        proposal_code = proposal_id
    else:
        # It's a UUID, need to query to find proposal code
        items = await db_client.query_items(
            pk=f"USER#{user.get('user_id')}", index_name="GSI1"
        )

        proposal_item = None
        for item in items:
            if item.get("id") == proposal_id:
                proposal_item = item
                break

        if not proposal_item:
            raise HTTPException(status_code=404, detail="Proposal not found")

        pk = proposal_item["PK"]
        proposal_code = proposal_item.get("proposalCode", proposal_id)

    # Load proposal metadata
    proposal = await db_client.get_item(pk=pk, sk="METADATA")

    if not proposal:
        raise HTTPException(status_code=404, detail="Proposal not found")

    # Verify ownership
    if proposal.get("user_id") != user.get("user_id"):
        raise HTTPException(status_code=403, detail="Access denied")

    return pk, proposal_code, proposal


# ==================== STEP 3: STRUCTURE & WORKPLAN ====================


@router.post("/{proposal_id}/analyze-step-3")
async def analyze_step_3(proposal_id: str, user=Depends(get_current_user)):
    """
    Step 3: Structure and Workplan Analysis

    Prerequisites: Step 1 (RFP) and Step 2 (Concept Evaluation) must be completed

    This endpoint:
    1. Verifies RFP analysis and concept evaluation are completed
    2. Generates proposal structure and workplan
    3. Returns analysis result

    Returns structure with:
    - narrative_overview: Text summary
    - proposal_mandatory: Required sections (Executive Summary, Problem Context, Proposed Approach)
    - proposal_outline: Full proposal structure
    - hcd_notes: Human-centered design notes
    """
    try:
        # Verify proposal ownership using helper
        pk, proposal_code, proposal = await _verify_proposal_access(proposal_id, user)

        # Check prerequisites
        if not proposal.get("rfp_analysis"):
            raise HTTPException(
                status_code=400,
                detail="Step 1 (RFP analysis) must be completed before Step 3.",
            )

        # Check for concept_evaluation or concept_analysis
        concept_eval = proposal.get("concept_evaluation") or proposal.get(
            "concept_analysis"
        )
        if not concept_eval:
            raise HTTPException(
                status_code=400,
                detail="Step 2 (Concept evaluation) must be completed before Step 3.",
            )

        print(f"✓ Prerequisites met for {proposal_code}")

        # Check if already completed
        if proposal.get("structure_workplan_analysis"):
            return {
                "status": "completed",
                "message": "Structure and workplan already analyzed",
                "data": proposal.get("structure_workplan_analysis"),
                "cached": True,
            }

        # Check if already processing
        if proposal.get("analysis_status_structure_workplan") == "processing":
            return {
                "status": "processing",
                "message": "Structure and workplan analysis already in progress",
                "started_at": proposal.get("structure_workplan_started_at"),
            }

        # Update status to processing BEFORE invoking worker
        await db_client.update_item(
            pk=pk,
            sk="METADATA",
            update_expression="SET analysis_status_structure_workplan = :status, structure_workplan_started_at = :started",
            expression_attribute_values={
                ":status": "processing",
                ":started": datetime.utcnow().isoformat(),
            },
        )

        # Invoke Worker Lambda asynchronously
        print(
            f"🚀 Invoking AnalysisWorkerFunction for structure workplan: {proposal_code}"
        )

        worker_function_arn = os.environ.get("WORKER_FUNCTION_NAME")
        if not worker_function_arn:
            raise Exception("WORKER_FUNCTION_NAME environment variable not set")

        lambda_client.invoke(
            FunctionName=worker_function_arn,
            InvocationType="Event",
            Payload=json.dumps(
                {"proposal_id": proposal_code, "analysis_type": "structure_workplan"}
            ),
        )

        print(f"✅ Structure workplan analysis worker invoked successfully")

        return {
            "status": "processing",
            "message": "Structure and workplan analysis started. Poll for completion.",
            "started_at": datetime.utcnow().isoformat(),
        }

    except HTTPException:
        raise
    except Exception as e:
        print(f"❌ Error in Step 3 analysis: {e}")
        import traceback

        traceback.print_exc()
        raise HTTPException(
            status_code=500, detail=f"Failed to analyze Step 3: {str(e)}"
        )


@router.get("/{proposal_id}/structure-workplan-status")
async def get_structure_workplan_status(
    proposal_id: str, user=Depends(get_current_user)
):
    """Poll for Structure Workplan analysis completion status"""
    try:
        # Verify proposal ownership using helper
        pk, proposal_code, proposal = await _verify_proposal_access(proposal_id, user)

        status = proposal.get("analysis_status_structure_workplan", "not_started")

        response = {
            "status": status,
            "started_at": proposal.get("structure_workplan_started_at"),
            "completed_at": proposal.get("structure_workplan_completed_at"),
            "error": proposal.get("structure_workplan_error"),
        }

        if status == "completed":
            response["data"] = proposal.get("structure_workplan_analysis")

        return response

    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to check status: {str(e)}")


@router.post("/{proposal_id}/generate-proposal-template")
async def generate_proposal_template(
    proposal_id: str, request: TemplateGenerationRequest, user=Depends(get_current_user)
):
    """
    Generate Word template from structure and workplan analysis.

    Prerequisites: Step 3 (Structure and Workplan) must be completed

    Request body:
    - selected_sections: Optional list of section titles to include (None = all sections)
    - user_comments: Optional dict of user comments per section {section_title: comment}

    Returns Word document with:
    - Proposal title and metadata
    - Narrative overview
    - Selected sections with purpose, guidance, and questions
    - User comments (if provided) displayed in blue italic text
    - Space for writing content
    """
    try:
        # Verify proposal ownership
        if proposal_id.startswith("PROP-"):
            pk = f"PROPOSAL#{proposal_id}"
            proposal_code = proposal_id
        else:
            items = await db_client.query_items(
                pk=f"USER#{user.get('user_id')}", index_name="GSI1"
            )

            proposal_item = None
            for item in items:
                if item.get("id") == proposal_id:
                    proposal_item = item
                    break

            if not proposal_item:
                raise HTTPException(status_code=404, detail="Proposal not found")

            pk = proposal_item["PK"]
            proposal_code = proposal_item.get("proposalCode", proposal_id)

        proposal = await db_client.get_item(pk=pk, sk="METADATA")

        if not proposal:
            raise HTTPException(status_code=404, detail="Proposal not found")

        if proposal.get("user_id") != user.get("user_id"):
            raise HTTPException(status_code=403, detail="Access denied")

        # Check prerequisite
        if not proposal.get("structure_workplan_analysis"):
            raise HTTPException(
                status_code=400,
                detail="Structure and workplan analysis must be completed before generating template.",
            )

        print(f"✓ Generating template for {proposal_code}")
        if request.selected_sections:
            print(f"  Selected sections: {len(request.selected_sections)}")
        if request.user_comments:
            print(f"  User comments: {len(request.user_comments)}")

        # Generate template
        from fastapi.responses import StreamingResponse

        service = ProposalTemplateGenerator()
        buffer = service.generate_template(
            proposal_id=proposal_code,
            selected_sections=request.selected_sections,
            user_comments=request.user_comments,
        )

        return StreamingResponse(
            buffer,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={
                "Content-Disposition": f"attachment; filename=proposal_template_{proposal_code}.docx"
            },
        )

    except HTTPException:
        raise
    except Exception as e:
        print(f"❌ Error generating template: {e}")
        import traceback

        traceback.print_exc()
        raise HTTPException(
            status_code=500, detail=f"Failed to generate template: {str(e)}"
        )


@router.post("/{proposal_id}/generate-ai-proposal-template")
async def generate_ai_proposal_template(
    proposal_id: str, request: TemplateGenerationRequest, user=Depends(get_current_user)
):
    """
    Generate AI-powered draft proposal based on all previous analyses.

    This endpoint:
    1. Validates all prerequisites (RFP, structure workplan, concept document)
    2. Sets status to "processing"
    3. Invokes Worker Lambda asynchronously
    4. Returns immediately with status

    Frontend should poll /proposal-template-status for completion

    Prerequisites:
    - RFP analysis must be completed
    - Structure workplan must be completed
    - Concept document must be generated

    Request body:
    - selected_sections: List of section titles to include (required)
    - user_comments: Optional dict of user comments per section
    """
    try:
        user_id = user.get("user_id")
        print(f"🟢 Generate AI proposal template request for proposal: {proposal_id}")

        # Verify proposal ownership using helper
        pk, proposal_code, proposal = await _verify_proposal_access(proposal_id, user)

        # Validate prerequisites
        if not proposal.get("rfp_analysis"):
            raise HTTPException(
                status_code=400, detail="RFP analysis must be completed first (Step 1)"
            )

        if not proposal.get("structure_workplan_analysis"):
            raise HTTPException(
                status_code=400,
                detail="Structure and workplan analysis must be completed first (Step 3)",
            )

        if not proposal.get("concept_document_v2"):
            raise HTTPException(
                status_code=400,
                detail="Concept document must be generated first (Step 2)",
            )

        # Validate selected sections
        selected_sections = request.selected_sections
        if not selected_sections or len(selected_sections) == 0:
            raise HTTPException(
                status_code=400, detail="At least one section must be selected"
            )

        print(f"   Selected sections: {len(selected_sections)}")
        if request.user_comments:
            print(f"   User comments: {len(request.user_comments)}")

        # Update status to processing
        await db_client.update_item(
            pk=pk,
            sk="METADATA",
            update_expression="SET proposal_template_status = :status, proposal_template_started_at = :started",
            expression_attribute_values={
                ":status": "processing",
                ":started": datetime.utcnow().isoformat(),
            },
        )

        # Invoke Worker Lambda asynchronously
        worker_function = os.getenv("WORKER_FUNCTION_NAME")

        payload = {
            "analysis_type": "proposal_template",
            "proposal_id": proposal_code,
            "user_id": user_id,
            "selected_sections": selected_sections,
            "user_comments": request.user_comments or {},
        }

        print(f"📡 Invoking worker lambda for proposal template generation: {proposal_code}")

        lambda_client.invoke(
            FunctionName=worker_function,
            InvocationType="Event",  # Asynchronous
            Payload=json.dumps(payload),
        )

        print(f"✅ Worker lambda invoked successfully for {proposal_code}")

        return {
            "status": "processing",
            "message": "Proposal template generation started. Poll /proposal-template-status for updates.",
        }

    except HTTPException:
        raise
    except Exception as e:
        print(f"❌ Error starting proposal template generation: {str(e)}")
        import traceback

        traceback.print_exc()
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/{proposal_id}/proposal-template-status")
async def get_proposal_template_status(proposal_id: str, user=Depends(get_current_user)):
    """
    Poll for AI proposal template generation completion status.

    Returns:
    - status: "not_started" | "processing" | "completed" | "failed"
    - started_at: ISO timestamp when generation started
    - completed_at: ISO timestamp when generation completed
    - error: Error message (if failed)
    - data: Generated proposal template content (if completed)
    """
    try:
        # Verify proposal ownership using helper
        pk, proposal_code, proposal = await _verify_proposal_access(proposal_id, user)

        status = proposal.get("proposal_template_status", "not_started")

        response = {
            "status": status,
            "started_at": proposal.get("proposal_template_started_at"),
            "completed_at": proposal.get("proposal_template_completed_at"),
            "error": proposal.get("proposal_template_error"),
        }

        if status == "completed":
            response["data"] = {
                "generated_proposal": proposal.get("proposal_template_content"),
                "sections": proposal.get("proposal_template_sections"),
                "metadata": proposal.get("proposal_template_metadata"),
                "s3_url": proposal.get("proposal_template_s3_url"),
            }

        return response

    except HTTPException:
        raise
    except Exception as e:
        print(f"❌ Error getting proposal template status: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to check status: {str(e)}")


@router.post("/{proposal_id}/use-generated-template-as-draft")
async def use_generated_template_as_draft(
    proposal_id: str, user=Depends(get_current_user)
):
    """
    Copy the AI-generated proposal template to the draft proposal location as DOCX.

    This endpoint:
    1. Verifies the generated template exists (proposal_template_content)
    2. Converts Markdown content to DOCX format
    3. Uploads to draft_proposal/ in S3
    4. Updates uploaded_files["draft-proposal"] in DynamoDB
    5. Sets draft_is_ai_generated flag

    Returns:
    - success: boolean
    - filename: Name of the created draft file
    - message: Status message
    """
    import io
    import re
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    def markdown_to_docx(markdown_content: str) -> bytes:
        """Convert markdown content to DOCX format."""
        doc = Document()

        # Set default font
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)

        lines = markdown_content.split('\n')
        i = 0

        while i < len(lines):
            line = lines[i].strip()

            # Skip empty lines
            if not line:
                i += 1
                continue

            # Headers
            if line.startswith('# '):
                p = doc.add_heading(line[2:], level=1)
            elif line.startswith('## '):
                p = doc.add_heading(line[3:], level=2)
            elif line.startswith('### '):
                p = doc.add_heading(line[4:], level=3)
            elif line.startswith('#### '):
                p = doc.add_heading(line[5:], level=4)
            # Bullet points
            elif line.startswith('- ') or line.startswith('* '):
                text = line[2:]
                # Handle bold and italic in bullet points
                text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)  # Remove bold markers
                text = re.sub(r'\*(.+?)\*', r'\1', text)  # Remove italic markers
                p = doc.add_paragraph(text, style='List Bullet')
            # Numbered lists
            elif re.match(r'^\d+\.\s', line):
                text = re.sub(r'^\d+\.\s', '', line)
                text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
                text = re.sub(r'\*(.+?)\*', r'\1', text)
                p = doc.add_paragraph(text, style='List Number')
            # Regular paragraphs
            else:
                # Handle bold text
                text = re.sub(r'\*\*(.+?)\*\*', r'\1', line)
                text = re.sub(r'\*(.+?)\*', r'\1', text)
                p = doc.add_paragraph(text)

            i += 1

        # Save to bytes
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()

    try:
        # Verify proposal ownership
        pk, proposal_code, proposal = await _verify_proposal_access(proposal_id, user)

        # Check if generated template exists
        template_content = proposal.get("proposal_template_content")
        if not template_content:
            raise HTTPException(
                status_code=400,
                detail="No generated proposal template found. Please generate a template first.",
            )

        # Get S3 bucket
        bucket = os.environ.get("PROPOSALS_BUCKET")
        if not bucket:
            raise HTTPException(
                status_code=500, detail="S3 bucket not configured"
            )

        # Convert markdown to DOCX
        docx_content = markdown_to_docx(template_content)

        # Define the draft filename (DOCX format)
        draft_filename = "ai_generated_draft.docx"
        draft_s3_key = f"{proposal_code}/documents/draft_proposal/{draft_filename}"

        # Upload DOCX content to draft location in S3
        s3_client = boto3.client("s3")
        s3_client.put_object(
            Bucket=bucket,
            Key=draft_s3_key,
            Body=docx_content,
            ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

        print(f"✅ Created AI draft DOCX at: s3://{bucket}/{draft_s3_key}")

        # Update DynamoDB with the new draft file
        current_uploaded_files = proposal.get("uploaded_files", {})
        current_uploaded_files["draft-proposal"] = [draft_filename]

        await db_client.update_item(
            pk=pk,
            sk="METADATA",
            update_expression="""
                SET uploaded_files = :files,
                    draft_is_ai_generated = :is_ai,
                    draft_source = :source,
                    updated_at = :updated
            """,
            expression_attribute_values={
                ":files": current_uploaded_files,
                ":is_ai": True,
                ":source": "ai_generated",
                ":updated": datetime.utcnow().isoformat(),
            },
        )

        print(f"✅ Updated DynamoDB with AI draft: {draft_filename}")

        return {
            "success": True,
            "filename": draft_filename,
            "message": "AI-generated template copied to draft location successfully",
            "s3_key": draft_s3_key,
        }

    except HTTPException:
        raise
    except Exception as e:
        print(f"❌ Error copying template to draft: {str(e)}")
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=str(e))


# ==================== STEP 4: DRAFT PROPOSAL FEEDBACK ====================


@router.post("/{proposal_id}/upload-draft-proposal")
async def upload_draft_proposal_file(
    proposal_id: str, file: UploadFile = File(...), user=Depends(get_current_user)
):
    """
    Upload draft proposal document to S3.

    Stores in {proposal_code}/documents/draft_proposal/{filename}
    Accepts PDF, DOC, and DOCX files up to 20MB.
    """
    from io import BytesIO

    try:
        # Validate file type
        allowed_extensions = (".pdf", ".doc", ".docx")
        if not file.filename or not file.filename.lower().endswith(allowed_extensions):
            raise HTTPException(
                status_code=400, detail="Only PDF, DOC, and DOCX files are supported"
            )

        # Read and validate file
        file_bytes = await file.read()
        file_size = len(file_bytes)

        if file_size == 0:
            raise HTTPException(status_code=400, detail="File is empty")
        if file_size > 20 * 1024 * 1024:  # 20MB limit
            raise HTTPException(
                status_code=400, detail="File size must be less than 20MB"
            )

        # Verify proposal ownership
        pk, proposal_code, proposal = await _verify_proposal_access(proposal_id, user)

        # Upload to S3
        s3_client = boto3.client("s3")
        bucket = os.environ.get("PROPOSALS_BUCKET")

        s3_key = f"{proposal_code}/documents/draft_proposal/{file.filename}"

        s3_client.put_object(
            Bucket=bucket,
            Key=s3_key,
            Body=BytesIO(file_bytes),
            ContentType=file.content_type or "application/octet-stream",
            Metadata={
                "proposal-id": proposal_id,
                "uploaded-by": user.get("user_id"),
                "original-size": str(file_size),
            },
        )

        print(f"✓ Draft proposal uploaded to S3: {s3_key}")

        # Update DynamoDB - store single file (replace previous)
        await db_client.update_item(
            pk=pk,
            sk="METADATA",
            update_expression="SET uploaded_files.#draft = :files, updated_at = :updated",
            expression_attribute_names={"#draft": "draft-proposal"},
            expression_attribute_values={
                ":files": [file.filename],
                ":updated": datetime.utcnow().isoformat(),
            },
        )

        return {
            "success": True,
            "message": "Draft proposal uploaded successfully",
            "filename": file.filename,
            "document_key": s3_key,
            "size": file_size,
        }

    except HTTPException:
        raise
    except Exception as e:
        print(f"❌ Error uploading draft proposal: {e}")
        import traceback

        traceback.print_exc()
        raise HTTPException(
            status_code=500, detail=f"Failed to upload draft proposal: {str(e)}"
        )


@router.delete("/{proposal_id}/documents/draft-proposal/{filename}")
async def delete_draft_proposal(
    proposal_id: str, filename: str, user=Depends(get_current_user)
):
    """
    Delete draft proposal document and clear feedback analysis.
    """
    try:
        # Verify proposal ownership
        pk, proposal_code, proposal = await _verify_proposal_access(proposal_id, user)

        # Delete from S3
        s3_client = boto3.client("s3")
        bucket = os.environ.get("PROPOSALS_BUCKET")
        s3_key = f"{proposal_code}/documents/draft_proposal/{filename}"

        try:
            s3_client.delete_object(Bucket=bucket, Key=s3_key)
            print(f"✓ Deleted draft proposal from S3: {s3_key}")
        except Exception as s3_error:
            print(f"⚠️ Warning: Could not delete from S3: {s3_error}")

        # Clear from DynamoDB (file reference and any feedback analysis)
        await db_client.update_item(
            pk=pk,
            sk="METADATA",
            update_expression="""
                SET uploaded_files.#draft = :empty,
                    updated_at = :updated
                REMOVE draft_feedback_analysis,
                       analysis_status_draft_feedback,
                       draft_feedback_started_at,
                       draft_feedback_completed_at,
                       draft_feedback_error
            """,
            expression_attribute_names={"#draft": "draft-proposal"},
            expression_attribute_values={
                ":empty": [],
                ":updated": datetime.utcnow().isoformat(),
            },
        )

        return {"success": True, "message": "Draft proposal deleted successfully"}

    except HTTPException:
        raise
    except Exception as e:
        print(f"❌ Error deleting draft proposal: {e}")
        raise HTTPException(
            status_code=500, detail=f"Failed to delete draft proposal: {str(e)}"
        )


class DraftFeedbackRequest(BaseModel):
    force: bool = False


@router.post("/{proposal_id}/analyze-draft-feedback")
async def analyze_draft_feedback(
    proposal_id: str,
    request: Optional[DraftFeedbackRequest] = None,
    user=Depends(get_current_user),
):
    """
    Start Draft Feedback analysis (Step 4).

    Analyzes the uploaded draft proposal against RFP requirements
    and provides section-by-section feedback.

    Prerequisites:
    - Step 1 (RFP Analysis) must be completed
    - Draft proposal document must be uploaded

    Args:
        force: If True, forces a new analysis even if one already exists.

    Returns immediately with status, poll /draft-feedback-status for completion.
    """
    try:
        pk, proposal_code, proposal = await _verify_proposal_access(proposal_id, user)

        # Get force flag from request body
        force = request.force if request else False

        # Check prerequisites
        if not proposal.get("rfp_analysis"):
            raise HTTPException(
                status_code=400,
                detail="Step 1 (RFP analysis) must be completed before draft feedback analysis.",
            )

        draft_docs = proposal.get("uploaded_files", {}).get("draft-proposal", [])
        if not draft_docs:
            raise HTTPException(
                status_code=400, detail="Please upload your draft proposal first."
            )

        # Check if already completed (unless force=True)
        if proposal.get("draft_feedback_analysis") and not force:
            return {
                "status": "completed",
                "message": "Draft feedback already analyzed",
                "data": proposal.get("draft_feedback_analysis"),
                "cached": True,
            }

        # If force=True, clear existing analysis data
        if force and proposal.get("draft_feedback_analysis"):
            print(
                f"🔄 Force re-analysis requested, clearing existing data for {proposal_code}"
            )
            await db_client.update_item(
                pk=pk,
                sk="METADATA",
                update_expression="""
                    REMOVE draft_feedback_analysis,
                           draft_feedback_completed_at,
                           draft_feedback_error
                    SET analysis_status_draft_feedback = :not_started,
                        updated_at = :updated
                """,
                expression_attribute_values={
                    ":not_started": "not_started",
                    ":updated": datetime.utcnow().isoformat(),
                },
            )

        # Check if already processing (unless force=True)
        if proposal.get("analysis_status_draft_feedback") == "processing" and not force:
            return {
                "status": "processing",
                "message": "Draft feedback analysis already in progress",
                "started_at": proposal.get("draft_feedback_started_at"),
            }

        # Update status to processing
        await db_client.update_item(
            pk=pk,
            sk="METADATA",
            update_expression="SET analysis_status_draft_feedback = :status, draft_feedback_started_at = :started",
            expression_attribute_values={
                ":status": "processing",
                ":started": datetime.utcnow().isoformat(),
            },
        )

        # Invoke Worker Lambda asynchronously
        worker_function_arn = os.environ.get("WORKER_FUNCTION_NAME")
        if not worker_function_arn:
            raise Exception("WORKER_FUNCTION_NAME environment variable not set")

        print(f"🚀 Invoking Lambda worker for draft feedback analysis: {proposal_code}")

        lambda_client.invoke(
            FunctionName=worker_function_arn,
            InvocationType="Event",  # Async invocation
            Payload=json.dumps(
                {"proposal_id": proposal_code, "analysis_type": "draft_feedback"}
            ),
        )

        return {
            "status": "processing",
            "message": "Draft feedback analysis started. Poll /draft-feedback-status for completion.",
            "started_at": datetime.utcnow().isoformat(),
        }

    except HTTPException:
        raise
    except Exception as e:
        print(f"❌ Error starting draft feedback analysis: {e}")
        import traceback

        traceback.print_exc()
        raise HTTPException(
            status_code=500, detail=f"Failed to start draft feedback analysis: {str(e)}"
        )


@router.get("/{proposal_id}/draft-feedback-status")
async def get_draft_feedback_status(proposal_id: str, user=Depends(get_current_user)):
    """
    Poll for Draft Feedback analysis completion status.

    Returns:
    - status: "not_started" | "processing" | "completed" | "failed"
    - data: Feedback analysis result (if completed)
    - error: Error message (if failed)
    """
    try:
        pk, proposal_code, proposal = await _verify_proposal_access(proposal_id, user)

        status = proposal.get("analysis_status_draft_feedback", "not_started")

        response = {
            "status": status,
            "started_at": proposal.get("draft_feedback_started_at"),
            "completed_at": proposal.get("draft_feedback_completed_at"),
            "error": proposal.get("draft_feedback_error"),
        }

        if status == "completed":
            response["data"] = proposal.get("draft_feedback_analysis")

        return response

    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(
            status_code=500, detail=f"Failed to check draft feedback status: {str(e)}"
        )


@router.post("/prompts/with-categories")
async def get_prompt_with_categories(
    request: PromptWithCategoriesRequest, user=Depends(get_current_user)
):
    """Get a prompt with categories injected as variables"""
    try:
        # Initialize prompt service
        prompt_service = PromptService()

        # Get prompt with injected categories
        prompt = await prompt_service.get_prompt_with_categories(
            request.prompt_id, request.categories
        )

        if not prompt:
            raise HTTPException(status_code=404, detail="Prompt not found")

        return {
            "prompt": prompt,
            "injected_categories": request.categories,
            "available_variables": [
                f"{{{{category_{i}}}}}" for i in range(1, len(request.categories) + 1)
            ]
            + ["{{categories}}"],
        }
    except Exception as e:
        raise HTTPException(
            status_code=500, detail=f"Failed to get prompt with categories: {str(e)}"
        )


@router.get("/{proposal_id}/download-draft")
async def download_draft_proposal(
    proposal_id: str, user=Depends(get_current_user)
):
    """
    Download the draft proposal file from S3.
    Returns the file as a downloadable response.
    """
    from fastapi.responses import Response

    try:
        # Verify proposal ownership
        pk, proposal_code, proposal = await _verify_proposal_access(proposal_id, user)

        # Get uploaded draft files
        uploaded_files = proposal.get("uploaded_files", {})
        draft_files = uploaded_files.get("draft-proposal", [])

        if not draft_files:
            raise HTTPException(
                status_code=404,
                detail="No draft proposal found. Please upload a draft first.",
            )

        # Get the first draft file
        draft_filename = draft_files[0]
        draft_s3_key = f"{proposal_code}/documents/draft_proposal/{draft_filename}"

        # Get S3 bucket
        bucket = os.environ.get("PROPOSALS_BUCKET")
        if not bucket:
            raise HTTPException(
                status_code=500, detail="S3 bucket not configured"
            )

        # Download file from S3
        s3_client = boto3.client("s3")
        try:
            response = s3_client.get_object(Bucket=bucket, Key=draft_s3_key)
            file_content = response["Body"].read()
            content_type = response.get("ContentType", "application/octet-stream")
        except Exception as e:
            raise HTTPException(
                status_code=404,
                detail=f"Draft file not found in storage: {str(e)}",
            )

        # Determine content type based on extension
        if draft_filename.lower().endswith(".pdf"):
            content_type = "application/pdf"
        elif draft_filename.lower().endswith(".docx"):
            content_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        elif draft_filename.lower().endswith(".doc"):
            content_type = "application/msword"

        # Return as downloadable file
        return Response(
            content=file_content,
            media_type=content_type,
            headers={
                "Content-Disposition": f'attachment; filename="{draft_filename}"'
            }
        )

    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(
            status_code=500, detail=f"Failed to download draft: {str(e)}"
        )


@router.get("/{proposal_id}/download-template-docx")
async def download_template_as_docx(
    proposal_id: str, user=Depends(get_current_user)
):
    """
    Download the AI-generated proposal template as a DOCX file.
    Returns the file as a downloadable response.
    """
    import io
    import re
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from fastapi.responses import Response

    def markdown_to_docx(markdown_content: str, title: str = "AI Generated Proposal Draft") -> bytes:
        """Convert markdown content to professionally formatted DOCX."""
        doc = Document()

        # Set default font
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)

        # Add title
        title_para = doc.add_heading(title, level=0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add generation date
        date_para = doc.add_paragraph(f"Generated on {datetime.now().strftime('%B %d, %Y')}")
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_run = date_para.runs[0]
        date_run.font.size = Pt(10)
        date_run.font.color.rgb = RGBColor(128, 128, 128)

        # Add separator
        doc.add_paragraph()

        lines = markdown_content.split('\n')
        i = 0

        while i < len(lines):
            line = lines[i].strip()

            # Skip empty lines
            if not line:
                i += 1
                continue

            # Headers
            if line.startswith('# '):
                p = doc.add_heading(line[2:], level=1)
            elif line.startswith('## '):
                p = doc.add_heading(line[3:], level=2)
            elif line.startswith('### '):
                p = doc.add_heading(line[4:], level=3)
            elif line.startswith('#### '):
                p = doc.add_heading(line[5:], level=4)
            # Bullet points
            elif line.startswith('- ') or line.startswith('* '):
                text = line[2:]
                # Handle bold and italic in bullet points
                text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)  # Remove bold markers
                text = re.sub(r'\*(.+?)\*', r'\1', text)  # Remove italic markers
                p = doc.add_paragraph(text, style='List Bullet')
            # Numbered lists
            elif re.match(r'^\d+\.\s', line):
                text = re.sub(r'^\d+\.\s', '', line)
                text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
                text = re.sub(r'\*(.+?)\*', r'\1', text)
                p = doc.add_paragraph(text, style='List Number')
            # Regular paragraphs
            else:
                # Handle bold text
                text = re.sub(r'\*\*(.+?)\*\*', r'\1', line)
                text = re.sub(r'\*(.+?)\*', r'\1', text)
                p = doc.add_paragraph(text)

            i += 1

        # Add footer
        doc.add_paragraph()
        footer_para = doc.add_paragraph("Generated by IGAD Proposal Writer - AI Assistant")
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_run = footer_para.runs[0]
        footer_run.font.size = Pt(9)
        footer_run.font.color.rgb = RGBColor(128, 128, 128)

        # Save to bytes
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()

    try:
        # Verify proposal ownership
        pk, proposal_code, proposal = await _verify_proposal_access(proposal_id, user)

        # Check if generated template exists
        template_content = proposal.get("proposal_template_content")
        if not template_content:
            raise HTTPException(
                status_code=400,
                detail="No generated proposal template found. Please generate a template first.",
            )

        # Get proposal title for the document
        proposal_title = proposal.get("title", "AI Generated Proposal Draft")

        # Convert markdown to DOCX
        docx_content = markdown_to_docx(template_content, proposal_title)

        # Generate filename
        filename = f"ai_proposal_draft_{proposal_code}_{datetime.now().strftime('%Y-%m-%d')}.docx"

        # Return as downloadable file
        return Response(
            content=docx_content,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={
                "Content-Disposition": f'attachment; filename="{filename}"'
            }
        )

    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(
            status_code=500, detail=f"Failed to generate DOCX: {str(e)}"
        )
