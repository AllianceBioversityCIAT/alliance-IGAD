"""
Concept Evaluation Service

Evaluates initial project concepts against RFP requirements to assess alignment:
- Thematic relevance to donor priorities
- Geographic and target population alignment
- Methodological approach suitability
- Completeness for proposal development

The service:
1. Retrieves concept document (text file or PDF/DOCX)
2. Loads RFP analysis for context
3. Combines with dynamic prompt from DynamoDB
4. Sends to Claude for structured evaluation
5. Returns alignment assessment and improvement recommendations
"""

import json
import os
import re
import time
import traceback
from io import BytesIO
from typing import Any, Dict, Optional

import boto3
from boto3.dynamodb.conditions import Attr
from docx import Document
from PyPDF2 import PdfReader

from app.database.client import db_client
from app.shared.ai.bedrock_service import BedrockService
from app.tools.proposal_writer.concept_evaluation.config import (
    CONCEPT_EVALUATION_SETTINGS,
)


class SimpleConceptAnalyzer:
    """
    Evaluates initial concepts against RFP requirements.

    Workflow:
    1. Load concept document from S3 (text, PDF, or DOCX)
    2. Retrieve RFP analysis for context
    3. Load evaluation prompt from DynamoDB
    4. Inject context into prompt
    5. Send to Claude via Bedrock
    6. Parse and return evaluation results
    """

    def __init__(self):
        """Initialize S3, DynamoDB, and Bedrock clients."""
        self.s3 = boto3.client("s3")
        self.bucket = os.environ.get("PROPOSALS_BUCKET")
        if not self.bucket:
            raise Exception("PROPOSALS_BUCKET environment variable not set")

        self.bedrock = BedrockService()
        self.dynamodb = boto3.resource("dynamodb")
        self.table_name = os.environ.get("TABLE_NAME", "igad-testing-main-table")

    def analyze_concept(
        self,
        proposal_id: str,
        rfp_analysis: Dict,
        reference_proposals_analysis: Optional[Dict] = None,
        existing_work_analysis: Optional[Dict] = None,
    ) -> Dict[str, Any]:
        """
        Evaluate concept against RFP requirements, reference proposals, and existing work.

        Args:
            proposal_id: Unique proposal identifier
            rfp_analysis: RFP analysis data from step 1
            reference_proposals_analysis: Reference proposals analysis from step 2 (optional)
            existing_work_analysis: Existing work analysis from step 2 (optional)

        Returns:
            Dict with structure:
            {
                "concept_analysis": {
                    "fit_assessment": {...},
                    "strong_aspects": [...],
                    "sections_needing_elaboration": [...],
                    "strategic_verdict": "..."
                },
                "status": "completed"
            }

        Raises:
            Exception: If concept not found or analysis fails
        """
        try:
            # Step 1: Retrieve proposal metadata
            print(f"📋 Loading proposal: {proposal_id}")
            proposal = db_client.get_item_sync(
                pk=f"PROPOSAL#{proposal_id}", sk="METADATA"
            )

            if not proposal:
                raise Exception(f"Proposal {proposal_id} not found")

            proposal_code = proposal.get("proposalCode")
            if not proposal_code:
                raise Exception("Proposal code not found")

            # Step 2: Retrieve concept document
            print(f"🔍 Locating concept document for: {proposal_code}")
            concept_text = self._get_concept_text_from_s3(proposal_code)

            # Step 3: Validate concept
            if not concept_text or len(concept_text) < 50:
                raise Exception(
                    "Concept text is too short or empty (minimum 50 characters)"
                )

            print(f"✅ Concept loaded: {len(concept_text)} characters")

            # Step 4: Load evaluation prompt
            print("📝 Loading evaluation prompt...")
            prompt_parts = self.get_prompt_from_dynamodb()

            if not prompt_parts:
                raise Exception("No active prompt found for concept evaluation")

            print("✅ Using DynamoDB prompt")

            # Step 5: Prepare prompt with context
            prepared_concept = self._prepare_concept_text(concept_text)
            unwrapped_rfp = self._unwrap_rfp_analysis(rfp_analysis)
            unwrapped_ref_proposals = self._unwrap_analysis(
                reference_proposals_analysis, "reference_proposals_analysis"
            )
            unwrapped_existing_work = self._unwrap_analysis(
                existing_work_analysis, "existing_work_analysis"
            )

            final_user_prompt = self._build_user_prompt(
                prompt_parts,
                prepared_concept,
                unwrapped_rfp,
                unwrapped_ref_proposals,
                unwrapped_existing_work,
            )

            # Step 6: Call Bedrock for evaluation
            print("📡 Sending to Bedrock for concept evaluation...")
            start_time = time.time()

            ai_response = self.bedrock.invoke_claude(
                system_prompt=prompt_parts["system_prompt"],
                user_prompt=final_user_prompt,
                max_tokens=CONCEPT_EVALUATION_SETTINGS.get("max_tokens", 15000),
                temperature=CONCEPT_EVALUATION_SETTINGS.get("temperature", 0.2),
                model_id=CONCEPT_EVALUATION_SETTINGS["model"],
            )

            elapsed = time.time() - start_time
            print(f"⏱️ Bedrock response time: {elapsed:.2f} seconds")

            # Step 7: Parse response
            print("🔄 Parsing AI response...")
            result = self.parse_response(ai_response)

            print("✅ Concept evaluation completed successfully")

            return {"concept_analysis": result, "status": "completed"}

        except Exception as e:
            print(f"❌ Concept evaluation failed: {str(e)}")
            traceback.print_exc()
            raise Exception(f"Concept analysis failed: {str(e)}")

    # ==================== PRIVATE HELPER METHODS ====================

    def _get_concept_text_from_s3(self, proposal_code: str) -> str:
        """
        Retrieve concept document from S3.

        Tries in order:
        1. Plain text file (concept_text.txt)
        2. PDF/DOC/DOCX files in concept folder

        Args:
            proposal_code: Proposal code for path construction

        Returns:
            Extracted concept text

        Raises:
            Exception: If no concept document found
        """
        concept_folder = f"{proposal_code}/documents/initial_concept/"

        # Try text file first (fastest)
        try:
            text_key = f"{concept_folder}concept_text.txt"
            print(f"📄 Trying concept text file: {text_key}")
            obj = self.s3.get_object(Bucket=self.bucket, Key=text_key)
            concept_text = obj["Body"].read().decode("utf-8")
            print(f"✅ Loaded concept text: {len(concept_text)} characters")
            return concept_text
        except Exception as e:
            print(f"ℹ️  No concept_text.txt found: {str(e)}")

        # Try document files (PDF/DOCX)
        print("🔍 Searching for concept document files...")
        response = self.s3.list_objects_v2(Bucket=self.bucket, Prefix=concept_folder)

        if "Contents" not in response or len(response["Contents"]) == 0:
            raise Exception("No concept content found. Please upload a document.")

        # Find and extract from first document file
        concept_file = self._find_document_file(response["Contents"])
        if not concept_file:
            raise Exception("No document files found (PDF/DOC/DOCX).")

        print(f"📥 Extracting text from: {concept_file}")
        return self.extract_text_from_file(concept_file)

    def _find_document_file(self, contents: list) -> Optional[str]:
        """
        Find first document file in S3 contents.

        Searches for: PDF and DOCX files (excludes legacy .doc format).

        Args:
            contents: List of S3 object summaries

        Returns:
            Document file key if found, None otherwise
        """
        for obj in contents:
            key = obj["Key"]
            # Only support PDF and DOCX (legacy .doc format not supported)
            if key.lower().endswith((".pdf", ".docx")) and not key.endswith("/"):
                return key
        return None

    def _prepare_concept_text(self, concept_text: str) -> str:
        """
        Prepare concept text for injection into prompt.

        Uses intelligent truncation:
        - Keeps 60% from beginning (context/problem)
        - Keeps 40% from end (solution/conclusion)
        - Adds truncation notice

        Args:
            concept_text: Full concept document text

        Returns:
            Prepared text (possibly truncated)
        """
        max_chars = CONCEPT_EVALUATION_SETTINGS.get("max_chars", 100000)

        if len(concept_text) <= max_chars:
            print(f"📄 Concept within limit ({len(concept_text)} chars)")
            return concept_text

        print(f"✂️  Truncating concept from {len(concept_text)} to {max_chars} chars")

        chars_to_keep = max_chars - 200
        beginning_chars = int(chars_to_keep * 0.6)
        ending_chars = int(chars_to_keep * 0.4)

        beginning = concept_text[:beginning_chars]
        ending = concept_text[-ending_chars:]

        return (
            f"{beginning}\n\n"
            "[... Middle section truncated - "
            f"total document: {len(concept_text)} characters ...]\n\n"
            f"{ending}"
        )

    def _unwrap_rfp_analysis(self, rfp_analysis: Dict) -> Dict:
        """
        Unwrap nested RFP analysis structure.

        RFP analysis may come wrapped as:
        {'rfp_analysis': {...}, 'status': '...'}

        Args:
            rfp_analysis: RFP analysis data (possibly nested)

        Returns:
            Unwrapped RFP analysis dict
        """
        if not isinstance(rfp_analysis, dict):
            return {}

        return rfp_analysis.get("rfp_analysis", rfp_analysis)

    def _unwrap_analysis(self, analysis: Optional[Dict], key: str) -> Dict:
        """
        Unwrap nested analysis structure (for reference proposals or existing work).

        Analysis may come wrapped as:
        {key: {...}, 'status': '...'}

        Args:
            analysis: Analysis data (possibly nested or None)
            key: The key to unwrap (e.g., "reference_proposals_analysis")

        Returns:
            Unwrapped analysis dict, or empty dict if None
        """
        if not analysis or not isinstance(analysis, dict):
            return {}

        return analysis.get(key, analysis)

    def _build_user_prompt(
        self,
        prompt_parts: Dict,
        concept_text: str,
        rfp_analysis: Dict,
        reference_proposals_analysis: Dict,
        existing_work_analysis: Dict,
    ) -> str:
        """
        Build complete user prompt with injected context.

        Injects:
        - RFP summary ({{rfp_analysis.summary}})
        - RFP extracted data ({{rfp_analysis.extracted_data}})
        - Reference proposals analysis ({{reference_proposal_analysis}} or {{reference_proposals_analysis}})
        - Existing work analysis ({{existing_work_analysis}})
        - Concept text ({{initial_concept}})

        Args:
            prompt_parts: Prompt template parts from DynamoDB
            concept_text: Prepared concept text
            rfp_analysis: RFP analysis data
            reference_proposals_analysis: Reference proposals analysis from step 2
            existing_work_analysis: Existing work analysis from step 2

        Returns:
            Complete user prompt ready for Claude
        """
        user_prompt = prompt_parts["user_prompt"]

        # Prepare RFP data for injection
        rfp_summary_json = json.dumps(rfp_analysis.get("summary", {}), indent=2)
        rfp_extracted_json = json.dumps(
            rfp_analysis.get("extracted_data", {}), indent=2
        )

        # Prepare Step 2 analyses for injection
        reference_proposals_json = (
            json.dumps(reference_proposals_analysis, indent=2)
            if reference_proposals_analysis
            else "{}"
        )
        existing_work_json = (
            json.dumps(existing_work_analysis, indent=2)
            if existing_work_analysis
            else "{}"
        )

        # Inject all placeholders
        user_prompt = user_prompt.replace("{{rfp_analysis.summary}}", rfp_summary_json)
        user_prompt = user_prompt.replace(
            "{{rfp_analysis.extracted_data}}", rfp_extracted_json
        )
        # Handle both singular and plural forms for reference proposals
        user_prompt = user_prompt.replace(
            "{{reference_proposal_analysis}}", reference_proposals_json
        )
        user_prompt = user_prompt.replace(
            "{{reference_proposals_analysis}}", reference_proposals_json
        )
        user_prompt = user_prompt.replace(
            "{{existing_work_analysis}}", existing_work_json
        )
        user_prompt = user_prompt.replace("{{initial_concept}}", concept_text)

        # Append output format
        final_prompt = f"{user_prompt}\n\n{prompt_parts['output_format']}".strip()

        print(f"📝 Final prompt: {len(final_prompt)} characters")
        print("   - RFP analysis: ✅")
        print(
            f"   - Reference proposals: {'✅' if reference_proposals_analysis else '⚠️  (empty)'}"
        )
        print(
            f"   - Existing work: {'✅' if existing_work_analysis else '⚠️  (empty)'}"
        )
        return final_prompt

    # ==================== FILE EXTRACTION ====================

    def extract_text_from_file(self, s3_key: str) -> str:
        """
        Extract text from document file (PDF, DOC, or DOCX).

        Args:
            s3_key: S3 key of the file

        Returns:
            Extracted text

        Raises:
            Exception: If file type not supported or extraction fails
        """
        try:
            obj = self.s3.get_object(Bucket=self.bucket, Key=s3_key)
            file_bytes = obj["Body"].read()

            if s3_key.lower().endswith(".pdf"):
                return self.extract_text_from_pdf(file_bytes)
            elif s3_key.lower().endswith(".docx"):
                return self.extract_text_from_docx(file_bytes)
            elif s3_key.lower().endswith(".doc"):
                print(
                    "❌ Legacy .doc format not supported. Please upload as .docx or .pdf"
                )
                return None
            else:
                print(f"❌ Unsupported file type: {s3_key}")
                return None

        except Exception as e:
            print(f"❌ File extraction failed: {str(e)}")
            return None

    def extract_text_from_pdf(self, pdf_bytes: bytes) -> str:
        """
        Extract text from PDF bytes.

        Args:
            pdf_bytes: PDF file content as bytes

        Returns:
            Extracted text

        Raises:
            Exception: If PDF extraction fails
        """
        try:
            pdf_file = BytesIO(pdf_bytes)
            reader = PdfReader(pdf_file)
            text = ""

            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"

            return text.strip()

        except Exception as e:
            raise Exception(f"PDF extraction failed: {str(e)}")

    def extract_text_from_docx(self, docx_bytes: bytes) -> str:
        """
        Extract text from DOCX bytes.

        Extracts from:
        - Paragraphs
        - Tables (with pipe-separated columns)

        Args:
            docx_bytes: DOCX file content as bytes

        Returns:
            Extracted text, or message if empty

        Raises:
            Exception: If DOCX extraction fails
        """
        try:
            docx_file = BytesIO(docx_bytes)
            document = Document(docx_file)
            text = ""

            # Extract from paragraphs
            for paragraph in document.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + "\n"

            # Extract from tables
            for table in document.tables:
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells]
                    text += " | ".join(row_text) + "\n"

            return text.strip() if text.strip() else "[Empty DOCX file]"

        except Exception as e:
            raise Exception(f"DOCX extraction failed: {str(e)}")

    # ==================== PROMPT MANAGEMENT ====================

    def get_prompt_from_dynamodb(self) -> Optional[Dict[str, str]]:
        """
        Load evaluation prompt from DynamoDB.

        Searches for active prompt with criteria:
        - is_active: true
        - section: "proposal_writer"
        - sub_section: "step-1"
        - categories: contains "Initial Concept"

        Returns:
            Dict with 'system_prompt', 'user_prompt', 'output_format', or None
        """
        try:
            table = self.dynamodb.Table(self.table_name)
            filter_expr = (
                Attr("is_active").eq(True)
                & Attr("section").eq("proposal_writer")
                & Attr("sub_section").eq("step-1")
                & Attr("categories").contains("Initial Concept")
            )

            # Handle DynamoDB pagination
            items = []
            response = table.scan(FilterExpression=filter_expr)
            items.extend(response.get("Items", []))

            while "LastEvaluatedKey" in response:
                response = table.scan(
                    FilterExpression=filter_expr,
                    ExclusiveStartKey=response["LastEvaluatedKey"],
                )
                items.extend(response.get("Items", []))

            if not items:
                print("⚠️  No active prompts found in DynamoDB")
                return None

            prompt_item = items[0]
            print(f"✅ Loaded prompt: {prompt_item.get('name', 'Unnamed')}")

            return {
                "system_prompt": prompt_item.get("system_prompt", ""),
                "user_prompt": prompt_item.get("user_prompt_template", ""),
                "output_format": prompt_item.get("output_format", ""),
            }

        except Exception as e:
            print(f"⚠️  Failed to load prompt from DynamoDB: {str(e)}")
            return None

    # ==================== RESPONSE PARSING ====================

    def parse_response(self, response: str) -> Dict[str, Any]:
        """
        Parse Claude's response into structured JSON.

        Handles multiple response formats:
        1. Valid JSON object
        2. JSON in markdown code block (```json ... ```)
        3. JSON with markdown markers
        4. Falls back to error structure if parsing fails

        Args:
            response: Raw response from Claude

        Returns:
            Parsed JSON dict or error structure
        """
        try:
            response = response.strip()

            # Try to extract JSON from code block
            json_match = re.search(r"```json\s*(\{.*?\})\s*```", response, re.DOTALL)
            if json_match:
                response = json_match.group(1)
            else:
                # Try to find JSON object directly
                json_match = re.search(r"\{.*?\}", response, re.DOTALL)
                if json_match:
                    response = json_match.group(0)
                else:
                    # Remove markdown markers
                    response = (
                        response.removeprefix("```json")
                        .removeprefix("```")
                        .removesuffix("```")
                    )

            response = response.strip()
            parsed = json.loads(response)
            print("✅ Response parsed successfully")
            return parsed

        except json.JSONDecodeError as e:
            print(f"⚠️  JSON parsing failed: {str(e)}")
            return {
                "raw_response": response,
                "parse_error": str(e),
            }
